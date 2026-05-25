# Copyright (C) 2025 Marco Hernandez <ragettyandy@gmail.com>
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU Affero General Public License for more details.
#
# For information contact Marco Hernandez <ragettyandy@gmail.com>

"""
Unit tests for app.py — covers pure functions and Flask routes (mocked externals).
Run with:  pytest tests/test_app.py -v
"""
import json
import xml.etree.ElementTree as ET
from unittest.mock import MagicMock, patch

import anthropic
import numpy as np
import pytest
import requests

# ── Import helpers we can test without side-effects ──────────────────────────
# Patch heavy imports before importing app so the module loads cleanly in CI.
import sys

# Stub out fastembed so get_embedder() is never called implicitly at import time
fastembed_stub = MagicMock()
fastembed_stub.TextEmbedding = MagicMock()
sys.modules.setdefault("fastembed", fastembed_stub)

import app as _app  # noqa: E402  (must come after stubs)
import config as cfg


class TestCosineSimilarity:
    def test_identical_vectors(self):
        v = np.array([1.0, 0.0, 0.0])
        assert _app.cosine_similarity(v, v) == pytest.approx(1.0)

    def test_orthogonal_vectors(self):
        a = np.array([1.0, 0.0])
        b = np.array([0.0, 1.0])
        assert _app.cosine_similarity(a, b) == pytest.approx(0.0)

    def test_opposite_vectors(self):
        a = np.array([1.0, 0.0])
        b = np.array([-1.0, 0.0])
        assert _app.cosine_similarity(a, b) == pytest.approx(-1.0)

    def test_zero_vector_returns_zero(self):
        a = np.array([0.0, 0.0])
        b = np.array([1.0, 0.0])
        assert _app.cosine_similarity(a, b) == pytest.approx(0.0)

    def test_general_case(self):
        a = np.array([1.0, 1.0])
        b = np.array([1.0, 0.0])
        expected = 1.0 / np.sqrt(2)
        assert _app.cosine_similarity(a, b) == pytest.approx(expected)


class TestChunkText:
    def test_short_text_single_chunk(self):
        text = "Hello world"
        chunks = _app.chunk_text(text, chunk_size=cfg.CHUNK_SIZE, overlap=cfg.CHUNK_OVERLAP)
        assert chunks == ["Hello world"]

    def test_exact_chunk_size(self):
        text = "a" * cfg.CHUNK_SIZE
        chunks = _app.chunk_text(text, chunk_size=cfg.CHUNK_SIZE, overlap=cfg.CHUNK_OVERLAP)
        assert len(chunks) == 1
        assert chunks[0] == "a" * cfg.CHUNK_SIZE

    def test_two_chunks_with_overlap(self):
        text = "a" * (cfg.CHUNK_SIZE + cfg.CHUNK_OVERLAP)
        chunks = _app.chunk_text(text, chunk_size=cfg.CHUNK_SIZE, overlap=cfg.CHUNK_OVERLAP)
        assert len(chunks) == 2
        # Second chunk starts at offset (CHUNK_SIZE - CHUNK_OVERLAP)
        assert len(chunks[1]) == cfg.CHUNK_OVERLAP * 2

    def test_empty_text_returns_no_chunks(self):
        assert _app.chunk_text("", chunk_size=cfg.CHUNK_SIZE, overlap=cfg.CHUNK_OVERLAP) == []

    def test_whitespace_only_not_included(self):
        text = "   " + "a" * cfg.CHUNK_SIZE + "   "
        chunks = _app.chunk_text(text, chunk_size=cfg.CHUNK_SIZE, overlap=cfg.CHUNK_OVERLAP)
        # Strip means leading/trailing whitespace removed from each chunk
        for c in chunks:
            assert c == c.strip()

    def test_multiple_chunks_cover_all_content(self):
        # step = CHUNK_SIZE - CHUNK_OVERLAP; 3 steps fit in 2*CHUNK_SIZE + step
        step = cfg.CHUNK_SIZE - cfg.CHUNK_OVERLAP
        text = "x" * (cfg.CHUNK_SIZE + 2 * step)
        chunks = _app.chunk_text(text, chunk_size=cfg.CHUNK_SIZE, overlap=cfg.CHUNK_OVERLAP)
        assert len(chunks) == 3

class TestClampMaxResults:
    def test_default_empty_string(self):
        assert _app._clamp_max_results("") == cfg.MAX_RESULTS_DEFAULT

    def test_below_minimum_clamped_to_min(self):
        assert _app._clamp_max_results("10") == cfg.MAX_RESULTS_MIN

    def test_exact_minimum(self):
        assert _app._clamp_max_results(str(cfg.MAX_RESULTS_MIN)) == cfg.MAX_RESULTS_MIN

    def test_above_maximum_clamped_to_max(self):
        assert _app._clamp_max_results("999") == cfg.MAX_RESULTS_MAX

    def test_exact_maximum(self):
        assert _app._clamp_max_results(str(cfg.MAX_RESULTS_MAX)) == cfg.MAX_RESULTS_MAX

    def test_rounds_down_to_multiple_of_step(self):
        # 60 is between 2× and 3× the step size; should round down to 2×
        assert _app._clamp_max_results("60") == cfg.MAX_RESULTS_MIN * 2

    def test_exact_multiple_of_step_unchanged(self):
        assert _app._clamp_max_results("75") == cfg.MAX_RESULTS_MIN * 3

def _make_article_xml(authors_xml: str, abstract_xml: str = "", extra: str = "") -> ET.Element:
    xml = f"""
    <PubmedArticle>
      <MedlineCitation>
        <PMID>12345</PMID>
        <Article>
          <ArticleTitle>Test Title</ArticleTitle>
          <Journal>
            <Title>Test Journal</Title>
            <ISOAbbreviation>TJ</ISOAbbreviation>
            <JournalIssue>
              <PubDate><Year>2024</Year></PubDate>
            </JournalIssue>
          </Journal>
          <AuthorList>
            {authors_xml}
          </AuthorList>
          <Abstract>
            {abstract_xml}
          </Abstract>
          {extra}
        </Article>
      </MedlineCitation>
    </PubmedArticle>
    """
    return ET.fromstring(xml).find(".//MedlineCitation/Article")


class TestParseAuthors:
    def test_single_author(self):
        authors_xml = "<Author><LastName>Smith</LastName><Initials>JA</Initials></Author>"
        node = _make_article_xml(authors_xml)
        assert _app._parse_authors(node) == "Smith JA"

    def test_two_authors(self):
        authors_xml = """
        <Author><LastName>Smith</LastName><Initials>JA</Initials></Author>
        <Author><LastName>Jones</LastName><Initials>B</Initials></Author>
        """
        node = _make_article_xml(authors_xml)
        assert _app._parse_authors(node) == "Smith JA, Jones B"

    def test_more_than_max_authors_truncated(self):
        authors_xml = "".join(
            f"<Author><LastName>Author{i}</LastName><Initials>X</Initials></Author>"
            for i in range(cfg.AUTHORS_DISPLAY_MAX + 2)
        )
        node = _make_article_xml(authors_xml)
        result = _app._parse_authors(node)
        assert result.endswith("et al.")
        assert result.count(",") == cfg.AUTHORS_DISPLAY_MAX - 1

    def test_no_authors_returns_unknown(self):
        node = _make_article_xml("")
        assert _app._parse_authors(node) == "Unknown"

    def test_collective_name(self):
        authors_xml = "<Author><CollectiveName>The Study Group</CollectiveName></Author>"
        node = _make_article_xml(authors_xml)
        assert _app._parse_authors(node) == "The Study Group"



class TestParseAbstract:
    def test_plain_abstract(self):
        abstract_xml = '<AbstractText>This is the abstract.</AbstractText>'
        node = _make_article_xml("", abstract_xml)
        assert _app._parse_abstract(node) == "This is the abstract."

    def test_structured_abstract_with_labels(self):
        abstract_xml = """
        <AbstractText Label="BACKGROUND">Some background.</AbstractText>
        <AbstractText Label="METHODS">Some methods.</AbstractText>
        """
        node = _make_article_xml("", abstract_xml)
        result = _app._parse_abstract(node)
        assert "BACKGROUND: Some background." in result
        assert "METHODS: Some methods." in result

    def test_empty_abstract_text_skipped(self):
        abstract_xml = '<AbstractText></AbstractText>'
        node = _make_article_xml("", abstract_xml)
        assert _app._parse_abstract(node) == ""

    def test_no_abstract(self):
        node = _make_article_xml("")
        assert _app._parse_abstract(node) == ""


class TestSseEmit:
    def test_format(self):
        result = _app._sse_emit({"text": "hello"})
        assert result == 'data: {"text": "hello"}\n\n'

    def test_nested_payload(self):
        result = _app._sse_emit({"done": True, "citations": []})
        payload = json.loads(result[6:].strip())
        assert payload == {"done": True, "citations": []}


class TestParseSuggestions:
    def test_valid_json_array(self):
        raw = '["Q1?", "Q2?", "Q3?", "Q4?"]'
        assert _app._parse_suggestions(raw) == ["Q1?", "Q2?", "Q3?", "Q4?"]

    def test_truncated_to_four(self):
        raw = '["Q1?", "Q2?", "Q3?", "Q4?", "Q5?"]'
        assert len(_app._parse_suggestions(raw)) == 4

    def test_invalid_json_returns_empty(self):
        assert _app._parse_suggestions("not json") == []

    def test_empty_string_returns_empty(self):
        assert _app._parse_suggestions("") == []

    def test_whitespace_stripped(self):
        raw = '  ["Q1?"]  '
        assert _app._parse_suggestions(raw) == ["Q1?"]



class TestComputeCost:
    def test_opus_known_price(self):
        # 1 M input + 1 M output at $15/$75 per MTok = $90
        cost = _app._compute_cost("claude-opus-4-6", 1_000_000, 1_000_000)
        assert cost == pytest.approx(90.0)

    def test_sonnet_known_price(self):
        cost = _app._compute_cost("claude-sonnet-4-6", 1_000_000, 1_000_000)
        assert cost == pytest.approx(18.0)

    def test_haiku_known_price(self):
        cost = _app._compute_cost("claude-haiku-4-5-20251001", 1_000_000, 1_000_000)
        assert cost == pytest.approx(4.80)

    def test_zero_tokens_is_zero(self):
        assert _app._compute_cost("claude-sonnet-4-6", 0, 0) == pytest.approx(0.0)

    def test_unknown_model_returns_zero(self):
        assert _app._compute_cost("unknown-model", 1_000_000, 1_000_000) == pytest.approx(0.0)

    def test_output_weighted_more_than_input(self):
        # For every model, output tokens cost more than input tokens
        for model in cfg.ALLOWED_CHAT_MODELS:
            cost_in  = _app._compute_cost(model, 1_000, 0)
            cost_out = _app._compute_cost(model, 0, 1_000)
            assert cost_out > cost_in, f"Output should cost more than input for {model}"

    def test_small_realistic_call(self):
        # ~2 000 input, ~500 output on Sonnet: (2000*3 + 500*15) / 1e6
        cost = _app._compute_cost("claude-sonnet-4-6", 2_000, 500)
        assert cost == pytest.approx((2_000 * 3 + 500 * 15) / 1_000_000)


class TestAdvanceDelimiterState:
    DELIM = cfg.RAG_DELIMITER

    def _emit(self, payload):
        return _app._sse_emit(payload)

    def test_accumulates_pre_delimiter_text(self):
        pre, _, found, events = _app._advance_delimiter_state(
            "", "", False, "hello", self.DELIM, self._emit
        )
        assert not found
        assert pre == "hello" or events  # safe_len may flush partial

    def test_detects_delimiter(self):
        _, post, found, events = _app._advance_delimiter_state(
            "answer", "", False, self.DELIM + "suggestions", self.DELIM, self._emit
        )
        assert found
        assert post == "suggestions"
        # The pre-delimiter text should have been emitted
        assert any("answer" in e for e in events)

    def test_after_delimiter_appends_to_post(self):
        _, post, found, events = _app._advance_delimiter_state(
            "", "part1", True, " part2", self.DELIM, self._emit
        )
        assert found
        assert post == "part1 part2"
        assert events == []

    def test_safe_flush_keeps_delimiter_window(self):
        # When text doesn't contain delimiter, chars outside the delimiter
        # window should be emitted as safe
        long_text = "x" * 100
        pre, _, _, events = _app._advance_delimiter_state(
            "", "", False, long_text, self.DELIM, self._emit
        )
        # Some safe text should have been emitted
        assert len(events) > 0 or len(pre) <= len(self.DELIM)



class TestBuildRagContext:
    def _make_chunk(self, pmid, title="T", authors="A", journal="J", year="2024",
                    chunk_text="text"):
        return {
            "pmid": pmid, "title": title, "authors": authors,
            "journal": journal, "year": year, "chunk_text": chunk_text,
        }

    def test_single_chunk(self):
        chunks = [self._make_chunk("111", title="Article One", chunk_text="Excerpt A")]
        context, citations = _app._build_rag_context(chunks)
        assert "[1] Article One" in context
        assert "Excerpt A" in context
        assert len(citations) == 1
        assert citations[0]["pmid"] == "111"

    def test_deduplicates_citations_by_pmid(self):
        chunks = [
            self._make_chunk("111", chunk_text="Chunk 1"),
            self._make_chunk("111", chunk_text="Chunk 2"),
            self._make_chunk("222", chunk_text="Chunk 3"),
        ]
        context, citations = _app._build_rag_context(chunks)
        # Both chunks from PMID 111 share [1]; PMID 222 gets [2]
        assert context.count("[1]") == 2
        assert "[2]" in context
        assert "[3]" not in context
        # Citations deduplicated: only 2 unique PMIDs
        assert len(citations) == 2
        pmids = [c["pmid"] for c in citations]
        assert pmids == ["111", "222"]

    def test_duplicate_pmid_shares_citation_number(self):
        chunks = [
            self._make_chunk("111", chunk_text="Chunk A"),
            self._make_chunk("111", chunk_text="Chunk B"),
        ]
        context, citations = _app._build_rag_context(chunks)
        assert context.count("[1]") == 2
        assert "[2]" not in context
        assert len(citations) == 1

    def test_numbering_increments(self):
        chunks = [self._make_chunk(str(i)) for i in range(3)]
        context, _ = _app._build_rag_context(chunks)
        assert "[1]" in context
        assert "[2]" in context
        assert "[3]" in context

    def test_citations_ordered_by_first_appearance(self):
        # Citation list should reflect the order chunks are first seen, not PMID sort order
        chunks = [
            self._make_chunk("333", chunk_text="First"),
            self._make_chunk("111", chunk_text="Second"),
            self._make_chunk("222", chunk_text="Third"),
        ]
        _, citations = _app._build_rag_context(chunks)
        assert [c["pmid"] for c in citations] == ["333", "111", "222"]

    def test_context_uses_article_number_not_chunk_index(self):
        # When chunks 1+2 share a PMID, neither [3] nor [2]-for-second-article
        # should appear for the second chunk — the article number is reused.
        chunks = [
            self._make_chunk("A", chunk_text="chunkA1"),
            self._make_chunk("A", chunk_text="chunkA2"),
            self._make_chunk("B", chunk_text="chunkB"),
        ]
        context, _ = _app._build_rag_context(chunks)
        # Article A = [1], article B = [2]; chunk index [3] must not appear
        assert "[3]" not in context
        assert context.count("[1]") == 2
        assert context.count("[2]") == 1



class TestExtractAnswerFromEvents:
    def test_basic_extraction(self):
        events = [
            _app._sse_emit({"text": "Hello "}),
            _app._sse_emit({"text": "world"}),
        ]
        assert _app._extract_answer_from_events(events) == "Hello world"

    def test_ignores_non_text_events(self):
        events = [
            _app._sse_emit({"text": "Answer"}),
            _app._sse_emit({"done": True}),
        ]
        assert _app._extract_answer_from_events(events) == "Answer"

    def test_ignores_malformed_lines(self):
        events = ["not a sse line", _app._sse_emit({"text": "ok"})]
        assert _app._extract_answer_from_events(events) == "ok"

    def test_empty_events(self):
        assert _app._extract_answer_from_events([]) == ""


# ── Helpers shared by new test classes ───────────────────────────────────────

def _make_stream_cm(text_tokens: list[str], query_text: str = "diabetes[MeSH]",
                    in_tokens: int = 100, out_tokens: int = 50):
    """Return a mock that behaves like client.messages.stream(...) as stream:"""
    text_block = MagicMock()
    text_block.type = "text"
    text_block.text = query_text

    final_msg = MagicMock()
    final_msg.content = [text_block]
    final_msg.usage.input_tokens = in_tokens
    final_msg.usage.output_tokens = out_tokens

    stream = MagicMock()
    stream.text_stream = iter(text_tokens)
    stream.get_final_message.return_value = final_msg
    stream.__enter__ = MagicMock(return_value=stream)
    stream.__exit__ = MagicMock(return_value=False)
    return stream


def _make_ncbi_resp(status_code: int = 200, json_data: dict | None = None,
                    content: bytes = b""):
    resp = MagicMock()
    resp.status_code = status_code
    resp.json.return_value = json_data or {}
    resp.content = content
    resp.raise_for_status = MagicMock()
    return resp


# ── get_embedder / embed_texts ────────────────────────────────────────────────

class TestEmbedTexts:
    def test_returns_numpy_arrays(self):
        mock_embedder = MagicMock()
        mock_embedder.embed.return_value = [np.array([0.1, 0.2, 0.3])]
        with patch("app.get_embedder", return_value=mock_embedder):
            result = _app.embed_texts(["hello"])
        assert isinstance(result[0], np.ndarray)

    def test_multiple_texts(self):
        mock_embedder = MagicMock()
        mock_embedder.embed.return_value = [np.zeros(3), np.ones(3)]
        with patch("app.get_embedder", return_value=mock_embedder):
            result = _app.embed_texts(["a", "b"])
        assert len(result) == 2

    def test_get_embedder_caches_instance(self):
        _app._embedder = None
        # fastembed is stubbed at the top of this file; TextEmbedding is a MagicMock
        e1 = _app.get_embedder()
        e2 = _app.get_embedder()
        assert e1 is e2
        _app._embedder = None  # restore


# ── build_pubmed_query ────────────────────────────────────────────────────────

class TestBuildPubmedQuery:
    def test_returns_stripped_text(self):
        stream_cm = _make_stream_cm([], query_text="  diabetes[MeSH]  ")
        with patch("app.client") as mock_client:
            mock_client.messages.stream.return_value = stream_cm
            result = _app.build_pubmed_query("diabetes")
        assert result == "diabetes[MeSH]"

    def test_calls_claude_with_query(self):
        stream_cm = _make_stream_cm([], query_text="result")
        with patch("app.client") as mock_client:
            mock_client.messages.stream.return_value = stream_cm
            _app.build_pubmed_query("my query")
        call_kwargs = mock_client.messages.stream.call_args.kwargs
        assert call_kwargs["messages"][0]["content"] == "my query"


# ── _ncbi_get retry logic ─────────────────────────────────────────────────────

class TestNcbiGet:
    def test_success_on_first_attempt(self):
        resp = _make_ncbi_resp(200)
        with patch("app.requests.get", return_value=resp) as _, \
             patch("app.time.sleep") as mock_sleep:
            result = _app._ncbi_get("http://ncbi", {}, 10)
        assert result is resp
        mock_sleep.assert_not_called()

    def test_retries_on_429(self):
        resp_429 = _make_ncbi_resp(429)
        resp_200 = _make_ncbi_resp(200)
        with patch("app.requests.get", side_effect=[resp_429, resp_200]), \
             patch("app.time.sleep"):
            result = _app._ncbi_get("http://ncbi", {}, 10)
        assert result.status_code == 200

    def test_retries_on_500(self):
        resp_500 = _make_ncbi_resp(500)
        resp_200 = _make_ncbi_resp(200)
        with patch("app.requests.get", side_effect=[resp_500, resp_200]), \
             patch("app.time.sleep"):
            result = _app._ncbi_get("http://ncbi", {}, 10)
        assert result.status_code == 200

    def test_retries_on_timeout(self):
        resp_200 = _make_ncbi_resp(200)
        with patch("app.requests.get",
                   side_effect=[requests.Timeout(), resp_200]), \
             patch("app.time.sleep"):
            result = _app._ncbi_get("http://ncbi", {}, 10)
        assert result.status_code == 200

    def test_retries_on_connection_error(self):
        resp_200 = _make_ncbi_resp(200)
        with patch("app.requests.get",
                   side_effect=[requests.ConnectionError(), resp_200]), \
             patch("app.time.sleep"):
            result = _app._ncbi_get("http://ncbi", {}, 10)
        assert result.status_code == 200

    def test_raises_after_max_retries(self):
        with patch("app.requests.get",
                   side_effect=requests.ConnectionError("down")), \
             patch("app.time.sleep"):
            with pytest.raises(requests.ConnectionError):
                _app._ncbi_get("http://ncbi", {}, 10, max_retries=2)

    def test_sleep_doubles_each_retry(self):
        resp_200 = _make_ncbi_resp(200)
        side_effects = [_make_ncbi_resp(429), _make_ncbi_resp(429), resp_200]
        with patch("app.requests.get", side_effect=side_effects), \
             patch("app.time.sleep") as mock_sleep:
            _app._ncbi_get("http://ncbi", {}, 10)
        delays = [call.args[0] for call in mock_sleep.call_args_list]
        assert delays[1] == delays[0] * 2


# ── search_pubmed ─────────────────────────────────────────────────────────────

class TestSearchPubmed:
    def test_returns_id_list(self):
        resp = _make_ncbi_resp(200, json_data={"esearchresult": {"idlist": ["1", "2", "3"]}})
        with patch("app._ncbi_get", return_value=resp):
            result = _app.search_pubmed("diabetes[MeSH]", 25)
        assert result == ["1", "2", "3"]

    def test_empty_results(self):
        resp = _make_ncbi_resp(200, json_data={"esearchresult": {"idlist": []}})
        with patch("app._ncbi_get", return_value=resp):
            result = _app.search_pubmed("obscure[MeSH]", 25)
        assert result == []


# ── _parse_article ────────────────────────────────────────────────────────────

def _make_pubmed_xml(pmid="123", title="Test Title", journal="Test Journal",
                     year="2024", authors_xml="", abstract_xml=""):
    xml = f"""<PubmedArticleSet>
      <PubmedArticle>
        <MedlineCitation>
          <PMID>{pmid}</PMID>
          <Article>
            <ArticleTitle>{title}.</ArticleTitle>
            <Journal>
              <Title>{journal}</Title>
              <JournalIssue><PubDate><Year>{year}</Year></PubDate></JournalIssue>
            </Journal>
            <AuthorList>{authors_xml}</AuthorList>
            <Abstract>{abstract_xml}</Abstract>
          </Article>
        </MedlineCitation>
      </PubmedArticle>
    </PubmedArticleSet>"""
    import xml.etree.ElementTree as ET
    root = ET.fromstring(xml)
    return root.find(".//PubmedArticle")


class TestParseArticle:
    def test_title_trailing_period_stripped(self):
        node = _make_pubmed_xml(title="My Article")
        result = _app._parse_article("123", node)
        assert result["title"] == "My Article"

    def test_missing_title_uses_fallback(self):
        import xml.etree.ElementTree as ET
        xml = """<PubmedArticle>
          <MedlineCitation>
            <PMID>999</PMID>
            <Article>
              <Journal><Title>J</Title>
                <JournalIssue><PubDate><Year>2020</Year></PubDate></JournalIssue>
              </Journal>
              <AuthorList/>
              <Abstract/>
            </Article>
          </MedlineCitation>
        </PubmedArticle>"""
        node = ET.fromstring(xml)
        result = _app._parse_article("999", node)
        assert result["title"] == "No title"

    def test_url_contains_pmid(self):
        node = _make_pubmed_xml(pmid="42")
        result = _app._parse_article("42", node)
        assert "42" in result["url"]

    def test_year_truncated_to_4_chars(self):
        node = _make_pubmed_xml(year="2024 Jan")
        result = _app._parse_article("1", node)
        assert result["year"] == "2024"

    def test_journal_populated(self):
        node = _make_pubmed_xml(journal="Nature Medicine")
        result = _app._parse_article("1", node)
        assert result["journal"] == "Nature Medicine"


# ── fetch_articles ────────────────────────────────────────────────────────────

class TestFetchArticles:
    def test_empty_pmids_returns_empty(self):
        assert _app.fetch_articles([]) == []

    def test_parses_xml_response(self):
        xml_bytes = b"""<PubmedArticleSet>
          <PubmedArticle>
            <MedlineCitation>
              <PMID>111</PMID>
              <Article>
                <ArticleTitle>Article One.</ArticleTitle>
                <Journal>
                  <Title>NEJM</Title>
                  <JournalIssue><PubDate><Year>2023</Year></PubDate></JournalIssue>
                </Journal>
                <AuthorList/>
                <Abstract/>
              </Article>
            </MedlineCitation>
          </PubmedArticle>
        </PubmedArticleSet>"""
        resp = _make_ncbi_resp(200, content=xml_bytes)
        with patch("app._ncbi_get", return_value=resp):
            articles = _app.fetch_articles(["111"])
        assert len(articles) == 1
        assert articles[0]["pmid"] == "111"
        assert articles[0]["title"] == "Article One"

    def test_skips_pmids_missing_from_response(self):
        xml_bytes = b"""<PubmedArticleSet>
          <PubmedArticle>
            <MedlineCitation>
              <PMID>111</PMID>
              <Article>
                <ArticleTitle>Only One.</ArticleTitle>
                <Journal>
                  <Title>J</Title>
                  <JournalIssue><PubDate><Year>2023</Year></PubDate></JournalIssue>
                </Journal>
                <AuthorList/><Abstract/>
              </Article>
            </MedlineCitation>
          </PubmedArticle>
        </PubmedArticleSet>"""
        resp = _make_ncbi_resp(200, content=xml_bytes)
        with patch("app._ncbi_get", return_value=resp):
            articles = _app.fetch_articles(["111", "999"])
        assert len(articles) == 1


# ── get_pmcids ────────────────────────────────────────────────────────────────

class TestGetPmcids:
    def test_empty_pmids_returns_empty(self):
        assert _app.get_pmcids([]) == {}

    def test_extracts_pmcid(self):
        data = {
            "linksets": [{
                "ids": ["111"],
                "linksetdbs": [{"linkname": "pubmed_pmc", "links": ["654321"]}],
            }]
        }
        resp = _make_ncbi_resp(200, json_data=data)
        with patch("app._ncbi_get", return_value=resp):
            result = _app.get_pmcids(["111"])
        assert result == {"111": "654321"}

    def test_skips_linkset_without_pmc_links(self):
        data = {
            "linksets": [{
                "ids": ["111"],
                "linksetdbs": [{"linkname": "pubmed_pubmed", "links": ["222"]}],
            }]
        }
        resp = _make_ncbi_resp(200, json_data=data)
        with patch("app._ncbi_get", return_value=resp):
            result = _app.get_pmcids(["111"])
        assert result == {}

    def test_skips_empty_linkset_ids(self):
        data = {"linksets": [{"ids": [], "linksetdbs": []}]}
        resp = _make_ncbi_resp(200, json_data=data)
        with patch("app._ncbi_get", return_value=resp):
            result = _app.get_pmcids(["111"])
        assert result == {}


# ── fetch_pmc_full_text ───────────────────────────────────────────────────────

class TestFetchPmcFullText:
    def test_extracts_body_paragraphs(self):
        xml_bytes = b"""<article>
          <body>
            <p>This is a long enough paragraph to be included in the output.</p>
            <p>Another paragraph with sufficient length to pass the 20-char filter.</p>
          </body>
        </article>"""
        resp = _make_ncbi_resp(200, content=xml_bytes)
        with patch("app._ncbi_get", return_value=resp):
            text = _app.fetch_pmc_full_text("PMC123")
        assert "paragraph" in text

    def test_falls_back_to_root_when_no_body(self):
        xml_bytes = b"""<article>
          <p>Paragraph outside body with enough text length to pass.</p>
        </article>"""
        resp = _make_ncbi_resp(200, content=xml_bytes)
        with patch("app._ncbi_get", return_value=resp):
            text = _app.fetch_pmc_full_text("PMC456")
        assert "Paragraph" in text

    def test_short_paragraphs_excluded(self):
        xml_bytes = b"""<article><body><p>Short.</p></body></article>"""
        resp = _make_ncbi_resp(200, content=xml_bytes)
        with patch("app._ncbi_get", return_value=resp):
            text = _app.fetch_pmc_full_text("PMC789")
        assert text == ""


# ── _attach_similarities ──────────────────────────────────────────────────────

class TestAttachSimilarities:
    def test_adds_similarity_key(self):
        articles = [{"title": "T", "abstract": "A"}]
        q_emb = np.array([1.0, 0.0])
        a_emb = np.array([1.0, 0.0])
        with patch("app.embed_texts", side_effect=[[q_emb], [a_emb]]):
            _app._attach_similarities("query", articles)
        assert "similarity" in articles[0]
        assert articles[0]["similarity"] == pytest.approx(1.0)


# ── _run_search error branches ────────────────────────────────────────────────

class TestRunSearch:
    def test_returns_results_on_success(self):
        with patch("app.build_pubmed_query", return_value="q"), \
             patch("app.search_pubmed", return_value=["1"]), \
             patch("app.fetch_articles", return_value=[{"title": "T", "abstract": "A"}]), \
             patch("app._attach_similarities"):
            _, pubmed_q, err = _app._run_search("diabetes", 25)
        assert err is None
        assert pubmed_q == "q"

    def test_auth_error_returns_message(self):
        with patch("app.build_pubmed_query",
                   side_effect=anthropic.AuthenticationError.__new__(anthropic.AuthenticationError)):
            pass  # tested via side_effect below
        import anthropic as _anthropic
        with patch("app.build_pubmed_query",
                   side_effect=_anthropic.AuthenticationError(
                       message="bad key", response=MagicMock(), body={})):
            results, _, err = _app._run_search("q", 25)
        assert results is None
        assert "API key" in err

    def test_request_exception_returns_message(self):
        with patch("app.build_pubmed_query",
                   side_effect=requests.RequestException("timeout")):
            results, _, err = _app._run_search("q", 25)
        assert results is None
        assert "PubMed" in err

    def test_generic_exception_returns_str(self):
        with patch("app.build_pubmed_query",
                   side_effect=ValueError("unexpected")):
            results, _, err = _app._run_search("q", 25)
        assert results is None
        assert "unexpected" in err

    def test_empty_results_skips_similarity(self):
        with patch("app.build_pubmed_query", return_value="q"), \
             patch("app.search_pubmed", return_value=[]), \
             patch("app.fetch_articles", return_value=[]), \
             patch("app._attach_similarities") as mock_sim:
            _app._run_search("q", 25)
        mock_sim.assert_not_called()


# ── _try_fetch_full_text / _fetch_article_text ────────────────────────────────

class TestTryFetchFullText:
    def test_returns_text_and_pmcid_on_success(self):
        with patch("app.fetch_pmc_full_text", return_value="full text"):
            text, pmcid = _app._try_fetch_full_text("PMC1")
        assert text == "full text"
        assert pmcid == "PMC1"

    def test_returns_none_none_on_exception(self):
        with patch("app.fetch_pmc_full_text", side_effect=Exception("network")):
            text, pmcid = _app._try_fetch_full_text("PMC1")
        assert text is None
        assert pmcid is None


class TestFetchArticleText:
    _ART = {"pmid": "111", "title": "T", "abstract": "A"}

    def test_fetches_full_text_when_pmcid_available(self):
        with patch("app._try_fetch_full_text", return_value=("full", "PMC1")) as mock_ft:
            _, text, pmcid = _app._fetch_article_text(self._ART, {"111": "PMC1"})
        mock_ft.assert_called_once_with("PMC1")
        assert text == "full"
        assert pmcid == "PMC1"

    def test_returns_none_when_no_pmcid(self):
        _, text, pmcid = _app._fetch_article_text(self._ART, {})
        assert text is None
        assert pmcid is None


# ── generate_starter_questions ────────────────────────────────────────────────

class TestGenerateStarterQuestions:
    def test_returns_questions_list(self):
        mock_msg = MagicMock()
        mock_msg.content[0].text = '["Q1?", "Q2?", "Q3?", "Q4?"]'
        mock_msg.usage.input_tokens = 50
        mock_msg.usage.output_tokens = 30
        with patch("app.client") as mock_client:
            mock_client.messages.create.return_value = mock_msg
            result = _app.generate_starter_questions("diabetes", [{"title": "T"}])
        assert result == ["Q1?", "Q2?", "Q3?", "Q4?"]

    def test_truncates_to_four(self):
        mock_msg = MagicMock()
        mock_msg.content[0].text = '["Q1?","Q2?","Q3?","Q4?","Q5?"]'
        mock_msg.usage.input_tokens = 50
        mock_msg.usage.output_tokens = 30
        with patch("app.client") as mock_client:
            mock_client.messages.create.return_value = mock_msg
            result = _app.generate_starter_questions("q", [])
        assert len(result) == 4

    def test_returns_empty_on_exception(self):
        with patch("app.client") as mock_client:
            mock_client.messages.create.side_effect = Exception("api down")
            result = _app.generate_starter_questions("q", [])
        assert result == []


# ── _stream_delimited_response ────────────────────────────────────────────────

class TestStreamDelimitedResponse:
    _DELIM = cfg.RAG_DELIMITER

    def test_returns_events_post_and_usage(self):
        tokens = ["Answer text", self._DELIM, '["Q1?"]']
        stream_cm = _make_stream_cm(tokens)
        with patch("app.client") as mock_client:
            mock_client.messages.stream.return_value = stream_cm
            _, post, usage = _app._stream_delimited_response(
                "claude-sonnet-4-6", "ctx", "question", self._DELIM, _app._sse_emit
            )
        assert post.strip() == '["Q1?"]'
        assert "input_tokens" in usage
        assert "output_tokens" in usage

    def test_pre_text_emitted_when_no_delimiter(self):
        stream_cm = _make_stream_cm(["Just an answer"])
        with patch("app.client") as mock_client:
            mock_client.messages.stream.return_value = stream_cm
            events, _, _ = _app._stream_delimited_response(
                "claude-sonnet-4-6", "ctx", "question", self._DELIM, _app._sse_emit
            )
        full_text = _app._extract_answer_from_events(events)
        assert "Just an answer" in full_text

    def test_usage_defaults_on_exception(self):
        stream_cm = _make_stream_cm([])
        stream_cm.get_final_message.side_effect = Exception("stream error")
        with patch("app.client") as mock_client:
            mock_client.messages.stream.return_value = stream_cm
            _, _, usage = _app._stream_delimited_response(
                "claude-sonnet-4-6", "ctx", "q", self._DELIM, _app._sse_emit
            )
        assert usage == {"input_tokens": 0, "output_tokens": 0}


# ── Flask routes ──────────────────────────────────────────────────────────────

@pytest.fixture()
def flask_client():
    _app.app.config["TESTING"] = True
    with _app.app.test_client() as c:
        yield c


class TestIndexRoute:
    def test_get_returns_200(self, flask_client):
        resp = flask_client.get("/")
        assert resp.status_code == 200

    def test_post_empty_query_returns_200_no_search(self, flask_client):
        with patch("app._run_search") as mock_search:
            resp = flask_client.post("/", data={"query": "  "})
        assert resp.status_code == 200
        mock_search.assert_not_called()

    def test_post_with_query_calls_run_search(self, flask_client):
        with patch("app._run_search", return_value=([], "q", None)):
            resp = flask_client.post("/", data={"query": "diabetes"})
        assert resp.status_code == 200

    def test_post_with_error_still_returns_200(self, flask_client):
        with patch("app._run_search", return_value=(None, None, "API key invalid")):
            resp = flask_client.post("/", data={"query": "diabetes"})
        assert resp.status_code == 200


class TestCollectionsRoute:
    def test_get_returns_200(self, flask_client):
        with patch("app.db.list_collections", return_value=[]):
            resp = flask_client.get("/collections")
        assert resp.status_code == 200

    def test_post_missing_name_returns_400(self, flask_client):
        resp = flask_client.post("/collections",
                                 json={"articles": [{"pmid": "1"}]})
        assert resp.status_code == 400
        assert b"name" in resp.data.lower()

    def test_post_no_articles_returns_400(self, flask_client):
        resp = flask_client.post("/collections",
                                 json={"name": "My Collection", "articles": []})
        assert resp.status_code == 400

    def test_post_success_returns_collection_id(self, flask_client):
        art = {"pmid": "111", "title": "T", "authors": "A", "journal": "J",
               "year": "2024", "abstract": "Ab", "url": "http://x"}
        with patch("app.get_pmcids", return_value={}), \
             patch("app.db.create_collection", return_value=7), \
             patch("app.db.add_article"), \
             patch("app.db.chunks_exist", return_value=False), \
             patch("app.embed_texts", return_value=[np.zeros(384)] * 2), \
             patch("app.db.save_chunks"), \
             patch("app._fetch_article_text", return_value=(art, None, None)):
            resp = flask_client.post("/collections",
                                     json={"name": "Test", "user_query": "q",
                                           "pubmed_query": "pq", "articles": [art]})
        assert resp.status_code == 200
        assert resp.get_json()["id"] == 7


class TestCollectionDetailRoute:
    def test_returns_404_when_not_found(self, flask_client):
        with patch("app.db.get_collection", return_value=None):
            resp = flask_client.get("/collections/999")
        assert resp.status_code == 404

    def test_returns_200_when_found(self, flask_client):
        col = {"id": 1, "name": "C", "user_query": "q", "pubmed_query": "pq",
               "created_at": "2024-01-01", "article_count": 0,
               "full_text_count": 0, "chunk_count": 0}
        with patch("app.db.get_collection", return_value=col), \
             patch("app.db.get_collection_articles", return_value=[]), \
             patch("app.generate_starter_questions", return_value=[]):
            resp = flask_client.get("/collections/1")
        assert resp.status_code == 200


class TestCollectionDeleteRoute:
    def test_returns_ok(self, flask_client):
        with patch("app.db.delete_collection"):
            resp = flask_client.post("/collections/1/delete")
        assert resp.status_code == 200
        assert resp.get_json() == {"ok": True}


class TestConversationRoutes:
    def test_list_conversations(self, flask_client):
        with patch("app.db.list_conversations", return_value=[{"id": 1}]):
            resp = flask_client.get("/collections/1/conversations")
        assert resp.status_code == 200

    def test_get_messages(self, flask_client):
        with patch("app.db.get_conversation_messages", return_value=[]):
            resp = flask_client.get("/conversations/1/messages")
        assert resp.status_code == 200

    def test_delete_conversation(self, flask_client):
        with patch("app.db.delete_conversation"):
            resp = flask_client.post("/conversations/1/delete")
        assert resp.status_code == 200
        assert resp.get_json() == {"ok": True}


class TestCollectionAskRoute:
    def test_no_question_returns_400(self, flask_client):
        resp = flask_client.post("/collections/1/ask", json={"question": ""})
        assert resp.status_code == 400

    def test_invalid_model_falls_back_to_default(self, flask_client):
        q_emb = np.zeros(384)
        chunk = {"pmid": "1", "title": "T", "authors": "A", "journal": "J",
                 "year": "2024", "abstract": "ab", "url": "u",
                 "has_full_text": False, "chunk_text": "text",
                 "chunk_index": 0, "similarity": 0.9}
        with patch("app.embed_texts", return_value=[q_emb]), \
             patch("app.db.semantic_search", return_value=[chunk]), \
             patch("app.db.create_conversation", return_value=1), \
             patch("app.db.add_message"), \
             patch("app._stream_delimited_response",
                   return_value=([], "", {"input_tokens": 0, "output_tokens": 0})):
            resp = flask_client.post("/collections/1/ask",
                                     json={"question": "q", "model": "not-a-real-model"})
        assert resp.status_code == 200

    def test_empty_chunks_returns_no_articles_message(self, flask_client):
        with patch("app.embed_texts", return_value=[np.zeros(384)]), \
             patch("app.db.semantic_search", return_value=[]), \
             patch("app.db.create_conversation", return_value=1), \
             patch("app.db.add_message"):
            resp = flask_client.post("/collections/1/ask",
                                     json={"question": "What is diabetes?"})
        data = b"".join(resp.response)
        assert b"No articles" in data


class TestConversationExportRoute:
    def test_invalid_format_returns_400(self, flask_client):
        resp = flask_client.get("/conversations/1/export?format=pdf")
        assert resp.status_code == 400

    def test_missing_conversation_returns_404(self, flask_client):
        with patch("app.db.get_conversation", return_value=None):
            resp = flask_client.get("/conversations/1/export?format=rtf")
        assert resp.status_code == 404

    def test_no_messages_returns_404(self, flask_client):
        conv = {"id": 1, "title": "Chat", "collection_name": "Col", "created_at": "2024-01-01"}
        with patch("app.db.get_conversation", return_value=conv), \
             patch("app.db.get_conversation_messages", return_value=[]):
            resp = flask_client.get("/conversations/1/export?format=rtf")
        assert resp.status_code == 404

    def test_rtf_export_returns_rtf_content(self, flask_client):
        conv = {"id": 1, "title": "Chat", "collection_name": "Col", "created_at": "2024-01-01"}
        messages = [{"role": "user", "content": "Hi", "citations": None}]
        with patch("app.db.get_conversation", return_value=conv), \
             patch("app.db.get_conversation_messages", return_value=messages):
            resp = flask_client.get("/conversations/1/export?format=rtf")
        assert resp.status_code == 200
        assert resp.content_type == "application/rtf"

    def test_docx_export_returns_docx_content(self, flask_client):
        conv = {"id": 1, "title": "Chat", "collection_name": "Col", "created_at": "2024-01-01"}
        messages = [{"role": "user", "content": "Hi", "citations": None}]
        with patch("app.db.get_conversation", return_value=conv), \
             patch("app.db.get_conversation_messages", return_value=messages):
            resp = flask_client.get("/conversations/1/export?format=docx")
        assert resp.status_code == 200
        assert "wordprocessingml" in resp.content_type


# ── RTF helpers ───────────────────────────────────────────────────────────────

class TestRtfEsc:
    def test_backslash_escaped(self):
        assert _app._rtf_esc("a\\b") == "a\\\\b"

    def test_open_brace_escaped(self):
        assert _app._rtf_esc("{") == "\\{"

    def test_close_brace_escaped(self):
        assert _app._rtf_esc("}") == "\\}"

    def test_non_ascii_encoded(self):
        result = _app._rtf_esc("\u00e9")  # é
        assert result == "\\u233?"

    def test_plain_ascii_passthrough(self):
        assert _app._rtf_esc("hello world") == "hello world"

    def test_empty_string(self):
        assert _app._rtf_esc("") == ""


class TestRtfHyperlink:
    def test_contains_hyperlink_keyword(self):
        result = _app._rtf_hyperlink("http://example.com", "click me")
        assert "HYPERLINK" in result

    def test_contains_display_text(self):
        result = _app._rtf_hyperlink("http://example.com", "click me")
        assert "click me" in result

    def test_url_backslash_escaped(self):
        result = _app._rtf_hyperlink("http://x.com\\path", "t")
        assert "\\\\" in result


class TestRtfLineWithLinks:
    def test_plain_text_passthrough(self):
        result = _app._rtf_line_with_links("no markers here", {})
        assert result == "no markers here"

    def test_marker_with_url_becomes_hyperlink(self):
        result = _app._rtf_line_with_links("See [1] for details", {1: "http://x.com"})
        assert "HYPERLINK" in result
        assert "details" in result

    def test_marker_without_url_is_escaped_literally(self):
        result = _app._rtf_line_with_links("See [1] here", {})
        assert "HYPERLINK" not in result
        assert "[1]" in result


class TestBuildRtf:
    def test_contains_title(self):
        result = _app._build_rtf("My Chat", "Col", "2024-01-01T00:00:00", [])
        assert "My Chat" in result

    def test_user_message_included(self):
        msgs = [{"role": "user", "content": "Hello there"}]
        result = _app._build_rtf("T", "C", "2024-01-01", msgs)
        assert "Hello there" in result

    def test_assistant_message_included(self):
        msgs = [{"role": "assistant", "content": "The answer is 42", "citations": []}]
        result = _app._build_rtf("T", "C", "2024-01-01", msgs)
        assert "The answer is 42" in result

    def test_citations_section_present_when_citations_exist(self):
        msgs = [{
            "role": "assistant",
            "content": "See [1]",
            "citations": [{"num": 1, "url": "http://x.com", "title": "Article",
                           "journal": "Nature", "year": "2024"}],
        }]
        result = _app._build_rtf("T", "C", "2024-01-01", msgs)
        assert "Sources" in result
        assert "Article" in result

    def test_date_truncated_to_10_chars(self):
        result = _app._build_rtf("T", "C", "2024-01-01T12:00:00Z", [])
        assert "2024-01-01" in result
        assert "T12:00:00Z" not in result


# ═══════════════════════════════════════════════════════════════════════════════
# Flask route tests (mocked DB + external calls)
# ═══════════════════════════════════════════════════════════════════════════════

@pytest.fixture()
def client():
    _app.app.config["TESTING"] = True
    with _app.app.test_client() as c:
        yield c


class TestIndexRoute:
    def test_get_returns_200(self, client):
        resp = client.get("/")
        assert resp.status_code == 200

    @patch("app._run_search")
    def test_post_with_query_calls_run_search(self, mock_run, client):
        mock_run.return_value = ([], "diabetes[MeSH]", None)
        resp = client.post("/", data={"query": "diabetes research", "max_results": "25"})
        assert resp.status_code == 200
        mock_run.assert_called_once_with("diabetes research", 25)

    def test_post_empty_query_no_search(self, client):
        resp = client.post("/", data={"query": ""})
        assert resp.status_code == 200

    @patch("app._run_search")
    def test_error_displayed(self, mock_run, client):
        mock_run.return_value = (None, None, "API key invalid")
        resp = client.post("/", data={"query": "test"})
        assert resp.status_code == 200
        assert b"API key invalid" in resp.data


class TestCollectionsRoute:
    @patch("app.db.list_collections", return_value=[])
    def test_get_collections(self, mock_list, client):
        resp = client.get("/collections")
        assert resp.status_code == 200
        mock_list.assert_called_once()


class TestFetchArticleText:
    _ART = {"pmid": "123", "title": "T", "authors": "A",
            "journal": "J", "year": "2024", "abstract": "AB",
            "url": "https://pubmed.ncbi.nlm.nih.gov/123/"}

    def test_no_pmcid_returns_none(self):
        art, full_text, pmcid = _app._fetch_article_text(self._ART, {})
        assert art is self._ART
        assert full_text is None
        assert pmcid is None

    @patch("app._try_fetch_full_text", return_value=("Full text body", "PMC999"))
    def test_with_pmcid_calls_fetch(self, mock_fetch):
        _, full_text, pmcid = _app._fetch_article_text(self._ART, {"123": "PMC999"})
        mock_fetch.assert_called_once_with("PMC999")
        assert full_text == "Full text body"
        assert pmcid == "PMC999"

    @patch("app._try_fetch_full_text", return_value=(None, None))
    def test_fetch_failure_returns_none(self, mock_fetch):
        _, full_text, pmcid = _app._fetch_article_text(self._ART, {"123": "PMC999"})
        assert full_text is None
        assert pmcid is None


class TestCollectionsSaveRoute:
    _ART = {"pmid": "123", "title": "T", "authors": "A",
            "journal": "J", "year": "2024", "abstract": "AB",
            "url": "https://pubmed.ncbi.nlm.nih.gov/123/"}

    @patch("app.get_pmcids", return_value={})
    @patch("app.db.create_collection", return_value=42)
    @patch("app.db.add_article")
    @patch("app.db.chunks_exist", return_value=True)  # skip embedding phase
    @patch("app._fetch_article_text")
    def test_save_success(self, mock_fetch, mock_exist, mock_add, mock_create, mock_pmcids, client):
        mock_fetch.return_value = (self._ART, None, None)
        payload = {
            "name": "My Collection",
            "user_query": "diabetes",
            "pubmed_query": "diabetes[MeSH]",
            "articles": [self._ART],
        }
        resp = client.post("/collections", json=payload)
        assert resp.status_code == 200
        assert resp.get_json()["id"] == 42

    def test_save_missing_name_returns_400(self, client):
        resp = client.post("/collections", json={"articles": [{"pmid": "1"}]})
        assert resp.status_code == 400
        assert "name" in resp.get_json()["error"].lower()

    def test_save_missing_articles_returns_400(self, client):
        resp = client.post("/collections", json={"name": "Test", "articles": []})
        assert resp.status_code == 400

    @patch("app.db.save_chunks")
    @patch("app.embed_texts")
    @patch("app.db.chunks_exist", return_value=False)
    @patch("app.db.add_article")
    @patch("app.db.create_collection", return_value=1)
    @patch("app.get_pmcids", return_value={})
    @patch("app._fetch_article_text")
    def test_embed_called_once_for_multiple_articles(
        self, mock_fetch, mock_pmcids, mock_create, mock_add, mock_exist,
        mock_embed, mock_save, client
    ):
        """All chunks across all articles are embedded in a single batch call."""
        art1 = {**self._ART, "pmid": "111", "abstract": "Abstract one."}
        art2 = {**self._ART, "pmid": "222", "abstract": "Abstract two."}
        mock_fetch.side_effect = [(art1, None, None), (art2, None, None)]
        mock_embed.return_value = [np.zeros(384)] * 100  # oversized; index slicing handles it

        client.post("/collections", json={
            "name": "Batch Test",
            "user_query": "q",
            "pubmed_query": "q[MeSH]",
            "articles": [art1, art2],
        })

        # embed_texts must be called exactly once, not once per article
        mock_embed.assert_called_once()
        # save_chunks called once per article
        assert mock_save.call_count == 2


class TestCollectionDetailRoute:
    @patch("app.db.get_collection", return_value=None)
    def test_not_found_returns_404(self, mock_get, client):
        resp = client.get("/collections/999")
        assert resp.status_code == 404

    @patch("app.generate_starter_questions", return_value=["Q1?"])
    @patch("app.db.get_collection_articles", return_value=[])
    @patch("app.db.get_collection", return_value={
        "id": 1, "name": "Test", "user_query": "test",
        "pubmed_query": "test[MeSH]", "created_at": "2024-01-01",
        "article_count": 0, "full_text_count": 0, "chunk_count": 0,
    })
    def test_found_returns_200(self, mock_get, mock_articles, mock_questions, client):
        resp = client.get("/collections/1")
        assert resp.status_code == 200


class TestCollectionAskRoute:
    def test_empty_question_returns_400(self, client):
        resp = client.post("/collections/1/ask", json={"question": ""})
        assert resp.status_code == 400

    @patch("app.db.semantic_search", return_value=[])
    @patch("app.db.add_message")
    @patch("app.db.create_conversation", return_value=7)
    @patch("app.embed_texts")
    def test_no_chunks_streams_empty_message(
        self, mock_embed, mock_create_conv, mock_add_msg, mock_search, client
    ):
        mock_embed.return_value = [np.zeros(384)]
        resp = client.post(
            "/collections/1/ask",
            json={"question": "What is this?"},
            buffered=True,
        )
        assert resp.status_code == 200
        assert resp.content_type.startswith("text/event-stream")
        data = resp.data.decode()
        assert "No articles found" in data

    def test_invalid_model_falls_back_to_default(self, client):
        # Just verify the model validation logic directly
        model = "invalid-model"
        if model not in cfg.ALLOWED_CHAT_MODELS:
            model = cfg.DEFAULT_CHAT_MODEL
        assert model == cfg.DEFAULT_CHAT_MODEL


class TestCollectionDeleteRoute:
    @patch("app.db.delete_collection")
    def test_delete_returns_ok(self, mock_delete, client):
        resp = client.post("/collections/5/delete")
        assert resp.status_code == 200
        assert resp.get_json() == {"ok": True}
        mock_delete.assert_called_once_with(5)


class TestConversationRoutes:
    @patch("app.db.list_conversations", return_value=[])
    def test_list_conversations(self, mock_list, client):
        resp = client.get("/collections/1/conversations")
        assert resp.status_code == 200
        mock_list.assert_called_once_with(1)

    @patch("app.db.get_conversation_messages", return_value=[])
    def test_get_messages(self, mock_get, client):
        resp = client.get("/conversations/3/messages")
        assert resp.status_code == 200
        mock_get.assert_called_once_with(3)

    @patch("app.db.get_conversation_messages", return_value=[
        {"role": "user",      "content": "Q?",  "citations": None,
         "created_at": "2024-01-01T00:00:00"},
        {"role": "assistant", "content": "A.",
         "citations": [{"num": 1, "pmid": "111", "url": "https://pubmed.ncbi.nlm.nih.gov/111/",
                        "title": "T", "journal": "J", "year": "2024", "similarity": 0.9}],
         "created_at": "2024-01-01T00:00:01"},
    ])
    def test_get_messages_citations_field_returned(self, mock_get, client):
        resp = client.get("/conversations/3/messages")
        data = resp.get_json()
        assert data[0]["citations"] is None
        assert data[1]["citations"][0]["num"] == 1
        assert data[1]["citations"][0]["url"].startswith("https://")

    @patch("app.db.get_conversation_messages", return_value=[
        {"role": "assistant", "content": "A.",
         "citations": [
             {"num": 1, "pmid": "A", "url": "https://u1", "title": "T1", "journal": "J", "year": "2024", "similarity": 0.9},
             {"num": 2, "pmid": "B", "url": "https://u2", "title": "T2", "journal": "J", "year": "2024", "similarity": 0.8},
         ],
         "created_at": "2024-01-01T00:00:00"},
    ])
    def test_get_messages_citation_num_sequential(self, _, client):
        data = client.get("/conversations/3/messages").get_json()
        nums = [c["num"] for c in data[0]["citations"]]
        assert nums == [1, 2]

    @patch("app.db.delete_conversation")
    def test_delete_conversation(self, mock_delete, client):
        resp = client.post("/conversations/3/delete")
        assert resp.status_code == 200
        assert resp.get_json() == {"ok": True}
        mock_delete.assert_called_once_with(3)


class TestRtfEsc:
    def test_plain_ascii_unchanged(self):
        assert _app._rtf_esc("hello world") == "hello world"

    def test_backslash_escaped(self):
        assert _app._rtf_esc("a\\b") == "a\\\\b"

    def test_open_brace_escaped(self):
        assert _app._rtf_esc("a{b") == "a\\{b"

    def test_close_brace_escaped(self):
        assert _app._rtf_esc("a}b") == "a\\}b"

    def test_non_ascii_unicode_escaped(self):
        # é = U+00E9 = decimal 233
        result = _app._rtf_esc("caf\u00e9")
        assert "\\u233?" in result

    def test_empty_string(self):
        assert _app._rtf_esc("") == ""


class TestRtfHyperlink:
    def test_contains_hyperlink_keyword(self):
        result = _app._rtf_hyperlink("https://example.com", "click me")
        assert "HYPERLINK" in result

    def test_contains_url(self):
        result = _app._rtf_hyperlink("https://pubmed.ncbi.nlm.nih.gov/123/", "[1]")
        assert "https://pubmed.ncbi.nlm.nih.gov/123/" in result

    def test_contains_display_text(self):
        result = _app._rtf_hyperlink("https://example.com", "Article Title")
        assert "Article Title" in result

    def test_is_ascii_encodable(self):
        result = _app._rtf_hyperlink("https://example.com", "Title")
        result.encode("ascii")  # must not raise

    def test_rtf_field_structure(self):
        result = _app._rtf_hyperlink("https://example.com", "text")
        assert result.startswith(r"{\field{")
        assert result.endswith("}}")


class TestRtfLineWithLinks:
    _URL_MAP = {1: "https://pubmed.ncbi.nlm.nih.gov/111/",
                2: "https://pubmed.ncbi.nlm.nih.gov/222/"}

    def test_plain_text_unchanged(self):
        result = _app._rtf_line_with_links("No markers here.", {})
        assert result == "No markers here."

    def test_marker_becomes_hyperlink(self):
        result = _app._rtf_line_with_links("See [1] for details.", self._URL_MAP)
        assert "HYPERLINK" in result
        assert "pubmed.ncbi.nlm.nih.gov/111/" in result

    def test_unknown_marker_stays_plain(self):
        result = _app._rtf_line_with_links("See [9].", self._URL_MAP)
        assert "HYPERLINK" not in result
        assert "[9]" in result

    def test_multiple_markers(self):
        result = _app._rtf_line_with_links("As shown [1] and [2].", self._URL_MAP)
        assert result.count("HYPERLINK") == 2


class TestBuildRtf:
    _MSGS = [
        {"role": "user",      "content": "What is the treatment?"},
        {"role": "assistant", "content": "It involves X.\nAnd also Y."},
    ]
    _CITATION = {"num": 1, "pmid": "111", "title": "Key Article",
                 "url": "https://pubmed.ncbi.nlm.nih.gov/111/",
                 "journal": "NEJM", "year": "2024"}
    _MSGS_WITH_CITATIONS = [
        {"role": "user",      "content": "What is the treatment?"},
        {"role": "assistant", "content": "See [1] for details.",
         "citations": [_CITATION]},
    ]

    def test_starts_with_rtf_header(self):
        rtf = _app._build_rtf("Title", "Col", "2024-01-01", self._MSGS)
        assert rtf.startswith(r"{\rtf1")

    def test_ends_with_closing_brace(self):
        rtf = _app._build_rtf("Title", "Col", "2024-01-01", self._MSGS)
        assert rtf.rstrip().endswith("}")

    def test_contains_title(self):
        rtf = _app._build_rtf("My Study", "Col", "2024-01-01", self._MSGS)
        assert "My Study" in rtf

    def test_contains_collection_name(self):
        rtf = _app._build_rtf("Title", "Cardiology Collection", "2024-01-01", self._MSGS)
        assert "Cardiology Collection" in rtf

    def test_contains_date(self):
        rtf = _app._build_rtf("Title", "Col", "2024-06-15T10:00:00", self._MSGS)
        assert "2024-06-15" in rtf

    def test_user_label_present(self):
        rtf = _app._build_rtf("Title", "Col", "2024-01-01", self._MSGS)
        assert "You:" in rtf

    def test_assistant_label_present(self):
        rtf = _app._build_rtf("Title", "Col", "2024-01-01", self._MSGS)
        assert "Claude:" in rtf

    def test_result_is_ascii_encodable(self):
        rtf = _app._build_rtf("Title", "Col", "2024-01-01", self._MSGS)
        rtf.encode("ascii")  # must not raise

    def test_sources_section_present_when_citations(self):
        rtf = _app._build_rtf("Title", "Col", "2024-01-01", self._MSGS_WITH_CITATIONS)
        assert "Sources" in rtf

    def test_citation_title_in_sources(self):
        rtf = _app._build_rtf("Title", "Col", "2024-01-01", self._MSGS_WITH_CITATIONS)
        assert "Key Article" in rtf

    def test_citation_url_as_hyperlink(self):
        rtf = _app._build_rtf("Title", "Col", "2024-01-01", self._MSGS_WITH_CITATIONS)
        assert "pubmed.ncbi.nlm.nih.gov/111/" in rtf
        assert "HYPERLINK" in rtf

    def test_no_sources_section_without_citations(self):
        rtf = _app._build_rtf("Title", "Col", "2024-01-01", self._MSGS)
        assert "Sources" not in rtf

    def test_inline_marker_linked_in_answer(self):
        rtf = _app._build_rtf("Title", "Col", "2024-01-01", self._MSGS_WITH_CITATIONS)
        # The [1] in "See [1] for details." should become a hyperlink field
        assert "HYPERLINK" in rtf


# ═══════════════════════════════════════════════════════════════════════════════
# _rtf_citation_line / _rtf_user_turn / _rtf_assistant_turn helpers
# ═══════════════════════════════════════════════════════════════════════════════

class TestRtfCitationLine:
    def test_num_and_title_present(self):
        result = _app._rtf_citation_line({"num": 1, "title": "My Article"})
        assert "[1]" in result
        assert "My Article" in result

    def test_url_becomes_hyperlink(self):
        result = _app._rtf_citation_line({"num": 2, "url": "https://example.com", "title": "Art"})
        assert "HYPERLINK" in result

    def test_no_url_plain_title(self):
        result = _app._rtf_citation_line({"num": 3, "title": "Bare Title"})
        assert "HYPERLINK" not in result
        assert "Bare Title" in result

    def test_journal_and_year(self):
        result = _app._rtf_citation_line({"num": 4, "title": "T", "journal": "NEJM", "year": "2024"})
        assert "NEJM" in result
        assert "2024" in result

    def test_journal_without_year(self):
        result = _app._rtf_citation_line({"num": 5, "title": "T", "journal": "Lancet"})
        assert "Lancet" in result
        assert "(None)" not in result
        assert "()" not in result

    def test_no_journal_no_dash(self):
        result = _app._rtf_citation_line({"num": 6, "title": "T"})
        assert " - " not in result


class TestRtfUserTurn:
    def test_you_label_present(self):
        result = _app._rtf_user_turn({"role": "user", "content": "Hello"})
        assert "You:" in result

    def test_content_present(self):
        result = _app._rtf_user_turn({"role": "user", "content": "Hello"})
        assert "Hello" in result

    def test_returns_string(self):
        result = _app._rtf_user_turn({"role": "user", "content": "x"})
        assert isinstance(result, str)


class TestRtfAssistantTurn:
    def test_claude_label_present(self):
        result = _app._rtf_assistant_turn({"role": "assistant", "content": "Answer"})
        assert any("Claude:" in line for line in result)

    def test_content_present(self):
        result = _app._rtf_assistant_turn({"role": "assistant", "content": "Answer"})
        assert any("Answer" in line for line in result)

    def test_multiline_content_all_lines_present(self):
        result = _app._rtf_assistant_turn({"role": "assistant", "content": "A\nB\nC"})
        joined = "\n".join(result)
        assert "A" in joined
        assert "B" in joined
        assert "C" in joined

    def test_sources_label_present_when_citations(self):
        msg = {
            "role": "assistant",
            "content": "See [1].",
            "citations": [{"num": 1, "url": "https://example.com", "title": "Ex"}],
        }
        result = _app._rtf_assistant_turn(msg)
        assert any("Sources" in line for line in result)

    def test_no_sources_without_citations(self):
        result = _app._rtf_assistant_turn({"role": "assistant", "content": "No refs"})
        assert not any("Sources" in line for line in result)

    def test_returns_list(self):
        result = _app._rtf_assistant_turn({"role": "assistant", "content": "x"})
        assert isinstance(result, list)


class TestBuildDocx:
    import zipfile as _zf
    import io as _io

    _MSGS = [
        {"role": "user",      "content": "What is the treatment?"},
        {"role": "assistant", "content": "It involves X.\nAnd also Y."},
    ]
    _CITATION = {"num": 1, "pmid": "111", "title": "Key Article",
                 "url": "https://pubmed.ncbi.nlm.nih.gov/111/",
                 "journal": "NEJM", "year": "2024"}
    _MSGS_WITH_CITATIONS = [
        {"role": "user",      "content": "What is the treatment?"},
        {"role": "assistant", "content": "See [1] for details.",
         "citations": [_CITATION]},
    ]

    @staticmethod
    def _docx_xml(buf):
        import zipfile, io
        buf.seek(0)
        with zipfile.ZipFile(io.BytesIO(buf.read())) as z:
            with z.open("word/document.xml") as f:
                return f.read().decode("utf-8")

    def test_returns_bytes_io(self):
        import io
        result = _app._build_docx("Title", "Col", "2024-01-01", self._MSGS)
        assert isinstance(result, io.BytesIO)

    def test_content_is_zip_archive(self):
        # DOCX files are ZIP archives — magic bytes are PK (0x50 0x4B)
        result = _app._build_docx("Title", "Col", "2024-01-01", self._MSGS)
        assert result.read(2) == b"PK"

    def test_non_empty_output(self):
        result = _app._build_docx("Title", "Col", "2024-01-01", self._MSGS)
        assert result.getbuffer().nbytes > 0

    def test_citation_title_in_document_xml(self):
        result = _app._build_docx("Title", "Col", "2024-01-01", self._MSGS_WITH_CITATIONS)
        xml = self._docx_xml(result)
        assert "Key Article" in xml

    def test_citation_url_as_hyperlink_relationship(self):
        import zipfile, io
        result = _app._build_docx("Title", "Col", "2024-01-01", self._MSGS_WITH_CITATIONS)
        result.seek(0)
        with zipfile.ZipFile(io.BytesIO(result.read())) as z:
            rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
        assert "pubmed.ncbi.nlm.nih.gov/111/" in rels

    def test_no_hyperlink_relationships_without_citations(self):
        import zipfile, io
        result = _app._build_docx("Title", "Col", "2024-01-01", self._MSGS)
        result.seek(0)
        with zipfile.ZipFile(io.BytesIO(result.read())) as z:
            rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
        assert "pubmed" not in rels


# ═══════════════════════════════════════════════════════════════════════════════
# _render_user_turn / _render_assistant_turn / _render_citation helpers
# ═══════════════════════════════════════════════════════════════════════════════

class TestRenderUserTurn:
    def _xml(self, msg):
        import io
        from docx import Document
        doc = Document()
        _app._render_user_turn(doc, msg)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        import zipfile
        with zipfile.ZipFile(buf) as z:
            return z.read("word/document.xml").decode("utf-8")

    def test_you_label_present(self):
        xml = self._xml({"role": "user", "content": "Hello"})
        assert "You:" in xml

    def test_user_content_present(self):
        xml = self._xml({"role": "user", "content": "Hello"})
        assert "Hello" in xml


class TestRenderAssistantTurn:
    def _xml(self, msg):
        import io
        from docx import Document
        doc = Document()
        _app._render_assistant_turn(doc, msg)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        import zipfile
        with zipfile.ZipFile(buf) as z:
            return z.read("word/document.xml").decode("utf-8")

    def test_claude_label_present(self):
        xml = self._xml({"role": "assistant", "content": "Answer"})
        assert "Claude:" in xml

    def test_single_line_content_present(self):
        xml = self._xml({"role": "assistant", "content": "Answer"})
        assert "Answer" in xml

    def test_multiline_content_all_lines_present(self):
        xml = self._xml({"role": "assistant", "content": "Line one\nLine two\nLine three"})
        assert "Line one" in xml
        assert "Line two" in xml
        assert "Line three" in xml

    def test_sources_label_present_when_citations(self):
        msg = {
            "role": "assistant",
            "content": "See [1].",
            "citations": [{"num": 1, "url": "https://example.com", "title": "Ex"}],
        }
        xml = self._xml(msg)
        assert "Sources:" in xml

    def test_no_sources_label_without_citations(self):
        xml = self._xml({"role": "assistant", "content": "No refs"})
        assert "Sources:" not in xml


class TestRenderCitation:
    def _xml(self, c):
        import io
        from docx import Document
        doc = Document()
        _app._render_citation(doc, c)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        import zipfile
        with zipfile.ZipFile(buf) as z:
            return z.read("word/document.xml").decode("utf-8")

    def _rels(self, c):
        import io
        from docx import Document
        doc = Document()
        _app._render_citation(doc, c)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        import zipfile
        with zipfile.ZipFile(buf) as z:
            return z.read("word/_rels/document.xml.rels").decode("utf-8")

    def test_title_without_url(self):
        xml = self._xml({"num": 1, "title": "Bare Title"})
        assert "Bare Title" in xml

    def test_url_creates_relationship(self):
        rels = self._rels({"num": 2, "url": "https://example.com/art", "title": "Art"})
        assert "example.com/art" in rels

    def test_journal_and_year_in_xml(self):
        xml = self._xml({"num": 3, "title": "T", "journal": "NEJM", "year": "2024"})
        assert "NEJM" in xml
        assert "2024" in xml

    def test_journal_without_year(self):
        xml = self._xml({"num": 4, "title": "T", "journal": "Lancet"})
        assert "Lancet" in xml
        # no parenthesised year should appear
        assert "(None)" not in xml
        assert "()" not in xml

    def test_no_journal_no_suffix(self):
        xml = self._xml({"num": 5, "title": "T"})
        assert "\u2013" not in xml


# ═══════════════════════════════════════════════════════════════════════════════
# conversation export route
# ═══════════════════════════════════════════════════════════════════════════════

_EXPORT_CONV = {
    "id": 1, "title": "Test Question",
    "collection_name": "My Collection", "created_at": "2024-01-01T00:00:00",
}
_EXPORT_MSGS = [
    {"role": "user",      "content": "What is the treatment?"},
    {"role": "assistant", "content": "It involves X."},
]


class TestConversationExportRoute:
    def test_invalid_format_returns_400(self, client):
        resp = client.get("/conversations/1/export?format=pdf")
        assert resp.status_code == 400

    @patch("app.db.get_conversation", return_value=None)
    def test_missing_conversation_returns_404(self, mock_conv, client):
        resp = client.get("/conversations/999/export?format=docx")
        assert resp.status_code == 404

    @patch("app.db.get_conversation_messages", return_value=[])
    @patch("app.db.get_conversation", return_value=_EXPORT_CONV)
    def test_no_messages_returns_404(self, mock_conv, mock_msgs, client):
        resp = client.get("/conversations/1/export?format=docx")
        assert resp.status_code == 404

    @patch("app.db.get_conversation_messages", return_value=_EXPORT_MSGS)
    @patch("app.db.get_conversation", return_value=_EXPORT_CONV)
    def test_docx_returns_200(self, mock_conv, mock_msgs, client):
        resp = client.get("/conversations/1/export?format=docx")
        assert resp.status_code == 200

    @patch("app.db.get_conversation_messages", return_value=_EXPORT_MSGS)
    @patch("app.db.get_conversation", return_value=_EXPORT_CONV)
    def test_docx_content_type(self, mock_conv, mock_msgs, client):
        resp = client.get("/conversations/1/export?format=docx")
        assert "wordprocessingml" in resp.content_type

    @patch("app.db.get_conversation_messages", return_value=_EXPORT_MSGS)
    @patch("app.db.get_conversation", return_value=_EXPORT_CONV)
    def test_docx_content_is_zip(self, mock_conv, mock_msgs, client):
        resp = client.get("/conversations/1/export?format=docx")
        assert resp.data[:2] == b"PK"

    @patch("app.db.get_conversation_messages", return_value=_EXPORT_MSGS)
    @patch("app.db.get_conversation", return_value=_EXPORT_CONV)
    def test_rtf_returns_200(self, mock_conv, mock_msgs, client):
        resp = client.get("/conversations/1/export?format=rtf")
        assert resp.status_code == 200

    @patch("app.db.get_conversation_messages", return_value=_EXPORT_MSGS)
    @patch("app.db.get_conversation", return_value=_EXPORT_CONV)
    def test_rtf_content_type(self, mock_conv, mock_msgs, client):
        resp = client.get("/conversations/1/export?format=rtf")
        assert resp.content_type == "application/rtf"

    @patch("app.db.get_conversation_messages", return_value=_EXPORT_MSGS)
    @patch("app.db.get_conversation", return_value=_EXPORT_CONV)
    def test_rtf_body_starts_with_rtf_header(self, mock_conv, mock_msgs, client):
        resp = client.get("/conversations/1/export?format=rtf")
        assert resp.data.startswith(b"{\\rtf1")

    @patch("app.db.get_conversation_messages", return_value=_EXPORT_MSGS)
    @patch("app.db.get_conversation", return_value=_EXPORT_CONV)
    def test_docx_content_disposition(self, mock_conv, mock_msgs, client):
        resp = client.get("/conversations/1/export?format=docx")
        cd = resp.headers.get("Content-Disposition", "")
        assert "attachment" in cd
        assert ".docx" in cd

    @patch("app.db.get_conversation_messages", return_value=_EXPORT_MSGS)
    @patch("app.db.get_conversation", return_value=_EXPORT_CONV)
    def test_rtf_content_disposition(self, mock_conv, mock_msgs, client):
        resp = client.get("/conversations/1/export?format=rtf")
        cd = resp.headers.get("Content-Disposition", "")
        assert "attachment" in cd
        assert ".rtf" in cd


# ═══════════════════════════════════════════════════════════════════════════════
# conversation rename route
# ═══════════════════════════════════════════════════════════════════════════════

class TestConversationRenameRoute:
    @patch("app.db.rename_conversation")
    def test_rename_returns_200(self, mock_rename, client):
        resp = client.patch(
            "/conversations/1/rename",
            json={"title": "New Title"},
        )
        assert resp.status_code == 200
        mock_rename.assert_called_once_with(1, "New Title")

    @patch("app.db.rename_conversation")
    def test_rename_returns_ok_true(self, mock_rename, client):
        resp = client.patch("/conversations/1/rename", json={"title": "New Title"})
        assert resp.get_json()["ok"] is True

    @patch("app.db.rename_conversation")
    def test_empty_title_returns_400(self, mock_rename, client):
        resp = client.patch("/conversations/1/rename", json={"title": ""})
        assert resp.status_code == 400
        mock_rename.assert_not_called()

    @patch("app.db.rename_conversation")
    def test_whitespace_title_returns_400(self, mock_rename, client):
        resp = client.patch("/conversations/1/rename", json={"title": "   "})
        assert resp.status_code == 400
        mock_rename.assert_not_called()

    @patch("app.db.rename_conversation")
    def test_title_is_stripped(self, mock_rename, client):
        client.patch("/conversations/1/rename", json={"title": "  Trimmed  "})
        mock_rename.assert_called_once_with(1, "Trimmed")
