# Copyright (C) 2025 Marco Hernandez <ragettyandy@gmail.com>
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU Affero General Public License for more details.
#
# For information contact Marco Hernandez <ragettyandy@gmail.com>

import csv
import datetime
import io
import json
import logging
import logging.handlers
import os
import re
import secrets
import time
import xml.etree.ElementTree as ET
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

import litellm
import numpy as np
import requests
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from flask import Flask, Response, g, jsonify, make_response, redirect, render_template, request, stream_with_context, url_for
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_wtf import CSRFProtect
from dotenv import load_dotenv

import auth
import db
import config as cfg

load_dotenv()

# ── Logging setup ─────────────────────────────────────────────────────────────

def setup_logging() -> logging.Logger:
    level_name = os.environ.get("LOG_LEVEL", "INFO").upper()
    level = getattr(logging, level_name, logging.INFO)

    fmt = logging.Formatter(
        "%(asctime)s | %(levelname)-8s | %(name)-18s | %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )

    log_dir = Path(__file__).parent / "logs"
    log_dir.mkdir(exist_ok=True)

    file_handler = logging.handlers.RotatingFileHandler(
        log_dir / "app.log", maxBytes=10 * 1024 * 1024, backupCount=5, encoding="utf-8"
    )
    file_handler.setFormatter(fmt)

    console_handler = logging.StreamHandler()
    console_handler.setFormatter(fmt)

    root = logging.getLogger("pubmed")
    root.setLevel(level)
    if not root.handlers:          # avoid duplicate handlers on Flask reloader restart
        root.addHandler(file_handler)
        root.addHandler(console_handler)

    root.info("Logging initialised | level=%s | log=%s", level_name, log_dir / "app.log")
    return root

_log       = setup_logging()
_api_log   = logging.getLogger("pubmed.api")
_claude_log = logging.getLogger("pubmed.claude")
_embed_log = logging.getLogger("pubmed.embed")

_starter_questions_cache: dict[int, list[str]] = {}
_SSE_MIMETYPE = "text/event-stream"

app = Flask(__name__)
app.jinja_env.globals["static_version"] = cfg.STATIC_VERSION
app.secret_key = cfg.FLASK_SECRET_KEY
if not cfg.FLASK_SECRET_KEY:
    _log.warning(
        "FLASK_SECRET_KEY is not set — the session (and CSRF tokens riding on it) "
        "won't work. Set FLASK_SECRET_KEY in your .env before relying on this for anything real."
    )
csrf = CSRFProtect(app)

limiter = Limiter(key_func=get_remote_address, app=app, storage_uri="memory://")

# ── Available models (resolved once at first request) ─────────────────────────

_available_models_cache: dict[str, dict] | None = None


def _discover_ollama_models() -> dict[str, dict]:
    try:
        resp = requests.get(f"{cfg.OLLAMA_BASE_URL}/api/tags", timeout=2)
        if resp.status_code != 200:
            return {}
        return {
            f"ollama/{m['name']}": {
                "display":      m["name"],
                "provider":     "ollama",
                "requires_key": "",
                "input_price":  0.0,
                "output_price": 0.0,
            }
            for m in resp.json().get("models", [])
        }
    except Exception:
        return {}


def _discover_ollama_cloud_models() -> dict[str, dict]:
    """Ollama Cloud's hosted model catalog — not "pulled" models, a fixed list per account."""
    if not cfg.OLLAMA_API_KEY:
        return {}
    try:
        resp = requests.get(
            f"{cfg.OLLAMA_CLOUD_BASE_URL}/api/tags",
            headers={"Authorization": f"Bearer {cfg.OLLAMA_API_KEY}"},
            timeout=5,
        )
        if resp.status_code != 200:
            return {}
        return {
            f"ollama_chat/{m['name']}": {
                "display":      m["name"],
                "provider":     "ollama-cloud",
                "requires_key": "",
                "input_price":  0.0,
                "output_price": 0.0,
            }
            for m in resp.json().get("models", [])
        }
    except Exception:
        return {}


def available_chat_models() -> dict[str, dict]:
    """Static (env-key-gated) models, cached once, plus fresh Ollama discovery each call."""
    global _available_models_cache
    static_models = _available_models_cache
    if static_models is None:
        static_models = {
            mid: meta
            for mid, meta in cfg.CHAT_MODELS.items()
            if os.getenv(meta.get("requires_key", ""), "")
        }
        _available_models_cache = static_models
        _log.info("Available static chat models: %s", list(static_models))
    result = dict(static_models)
    result.update(_discover_ollama_models())
    result.update(_discover_ollama_cloud_models())
    return result


# ── Embeddings ────────────────────────────────────────────────────────────────

_embedder = None


def get_embedder():
    global _embedder
    if _embedder is None:
        from fastembed import TextEmbedding
        _embedder = TextEmbedding(cfg.EMBEDDING_MODEL, threads=cfg.EMBED_THREADS)
        _embed_log.info("Embedder loaded model=%s threads=%d", cfg.EMBEDDING_MODEL, cfg.EMBED_THREADS)
    return _embedder


def embed_texts(texts: list[str]) -> list[np.ndarray]:
    _embed_log.debug("Embedding %d text(s)", len(texts))
    t0 = time.perf_counter()
    result = [np.array(v, dtype=np.float32) for v in get_embedder().embed(texts, batch_size=cfg.EMBED_BATCH_SIZE)]
    elapsed = time.perf_counter() - t0
    dim = len(result[0]) if result else 0
    _embed_log.info("Embedded texts=%d dim=%d duration=%.3fs", len(texts), dim, elapsed)
    return result


def cosine_similarity(a: np.ndarray, b: np.ndarray) -> float:
    denom = np.linalg.norm(a) * np.linalg.norm(b)
    return float(np.dot(a, b) / denom) if denom > 0 else 0.0


# ── Text chunking ─────────────────────────────────────────────────────────────

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start += chunk_size - overlap
    return chunks


# ── PubMed helpers ────────────────────────────────────────────────────────────

def build_pubmed_query(english_query: str) -> str:
    _claude_log.debug("op=build_query query=%r", english_query[:120])
    t0 = time.perf_counter()
    response = litellm.completion(
        model=cfg.QUERY_BUILDER_MODEL,
        max_tokens=cfg.MAX_TOKENS_PUBMED_QUERY,
        thinking={"type": "enabled", "budget_tokens": cfg.PUBMED_QUERY_THINKING_BUDGET},
        messages=[
            {"role": "system", "content": cfg.PROMPT_PUBMED_QUERY},
            {"role": "user",   "content": english_query},
        ],
    )
    result = (response.choices[0].message.content or "").strip()
    elapsed = time.perf_counter() - t0
    usage = response.usage
    _claude_log.info(
        "op=build_query model=%s in_tokens=%d out_tokens=%d duration=%.2fs",
        cfg.QUERY_BUILDER_MODEL,
        getattr(usage, "prompt_tokens", 0),
        getattr(usage, "completion_tokens", 0),
        elapsed,
    )
    _claude_log.debug("op=build_query result=%r", result)
    return result


# ── NCBI request helper with retry / back-off ─────────────────────────────────

def _ncbi_get(url: str, params: dict, timeout: int, max_retries: int = 4) -> requests.Response:
    """GET with exponential back-off on transient NCBI errors (429, 5xx, timeouts).

    Delays: 1 s, 2 s, 4 s, 8 s (doubles each attempt, capped at 30 s).
    Raises the last exception if all retries are exhausted.
    """
    if cfg.PUBMED_API_KEY:
        params = {**params, "api_key": cfg.PUBMED_API_KEY}
    delay = 1.0
    for attempt in range(max_retries + 1):
        try:
            resp = requests.get(url, params=params, timeout=timeout)
            if resp.status_code == 429 or resp.status_code >= 500:
                raise requests.HTTPError(response=resp)
            return resp
        except (requests.Timeout, requests.ConnectionError, requests.HTTPError) as exc:
            _ncbi_handle_retry(url, attempt, max_retries, delay, exc)
            time.sleep(delay)
            delay = min(delay * 2, cfg.NCBI_BACKOFF_MAX)


def _ncbi_handle_retry(url: str, attempt: int, max_retries: int, delay: float, exc: Exception):
    """Log and raise on final attempt, or log a warning and allow retry."""
    if attempt == max_retries:
        _api_log.error(
            "op=ncbi_get url=%s attempt=%d/%d FAILED: %s",
            url, attempt + 1, max_retries + 1, exc,
        )
        raise exc
    _api_log.warning(
        "op=ncbi_get url=%s attempt=%d/%d retrying in %.0fs: %s",
        url, attempt + 1, max_retries + 1, delay, exc,
    )


def _ncbi_post(url: str, params: dict, timeout: int, max_retries: int = 4) -> requests.Response:
    """POST with exponential back-off — used for long queries that exceed GET URL limits."""
    if cfg.PUBMED_API_KEY:
        params = {**params, "api_key": cfg.PUBMED_API_KEY}
    delay = 1.0
    for attempt in range(max_retries + 1):
        try:
            resp = requests.post(url, data=params, timeout=timeout)
            if resp.status_code == 429 or resp.status_code >= 500:
                raise requests.HTTPError(response=resp)
            return resp
        except (requests.Timeout, requests.ConnectionError, requests.HTTPError) as exc:
            _ncbi_handle_retry(url, attempt, max_retries, delay, exc)
            time.sleep(delay)
            delay = min(delay * 2, cfg.NCBI_BACKOFF_MAX)


def search_pubmed(query: str, max_results: int = 25) -> list[str]:
    _api_log.debug("op=esearch query=%r max=%d", query[:120], max_results)
    t0 = time.perf_counter()
    resp = _ncbi_post(
        f"{cfg.PUBMED_BASE}/esearch.fcgi",
        params={"db": "pubmed", "term": query, "retmax": max_results, "retmode": "json", "sort": "relevance"},
        timeout=cfg.TIMEOUT_ESEARCH,
    )
    resp.raise_for_status()
    ids = resp.json()["esearchresult"]["idlist"]
    elapsed = time.perf_counter() - t0
    _api_log.info("op=esearch status=%d results=%d duration=%.2fs", resp.status_code, len(ids), elapsed)
    return ids


def _parse_authors(article_node) -> str:
    authors = []
    for author in article_node.findall(".//AuthorList/Author"):
        last = author.findtext("LastName") or ""
        initials = author.findtext("Initials") or ""
        collective = author.findtext("CollectiveName") or ""
        name = f"{last} {initials}".strip() if last else collective
        if name:
            authors.append(name)
    author_str = ", ".join(authors[:cfg.AUTHORS_DISPLAY_MAX])
    if len(authors) > cfg.AUTHORS_DISPLAY_MAX:
        author_str += " et al."
    return author_str or "Unknown"


def _parse_abstract(article_node) -> str:
    parts = []
    for ab in article_node.findall(".//Abstract/AbstractText"):
        text = (ab.text or "").strip()
        if not text:
            continue
        label = ab.get("Label")
        parts.append(f"{label}: {text}" if label else text)
    return "\n\n".join(parts)


def _parse_article(pmid: str, node) -> dict:
    article = node.find(".//MedlineCitation/Article")
    title = (article.findtext("ArticleTitle") or "No title").rstrip(".")
    journal = (
        article.findtext(".//Journal/Title")
        or article.findtext(".//Journal/ISOAbbreviation")
        or ""
    )
    year = (
        article.findtext(".//Journal/JournalIssue/PubDate/Year")
        or article.findtext(".//Journal/JournalIssue/PubDate/MedlineDate")
        or ""
    )[:4]
    return {
        "pmid": pmid,
        "title": title,
        "authors": _parse_authors(article),
        "journal": journal,
        "year": year,
        "abstract": _parse_abstract(article),
        "url": f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/",
    }


def fetch_articles(pmids: list[str]) -> list[dict]:
    if not pmids:
        return []
    _api_log.debug("op=efetch pmids=%d", len(pmids))
    t0 = time.perf_counter()
    resp = _ncbi_get(
        f"{cfg.PUBMED_BASE}/efetch.fcgi",
        params={"db": "pubmed", "id": ",".join(pmids), "rettype": "abstract", "retmode": "xml"},
        timeout=cfg.TIMEOUT_EFETCH,
    )
    resp.raise_for_status()
    _api_log.debug("op=efetch status=%d bytes=%d", resp.status_code, len(resp.content))
    root = ET.fromstring(resp.content)

    nodes_by_pmid = {
        node.findtext(".//MedlineCitation/PMID"): node
        for node in root.findall(".//PubmedArticle")
        if node.findtext(".//MedlineCitation/PMID")
    }

    articles = [
        _parse_article(pmid, nodes_by_pmid[pmid])
        for pmid in pmids
        if pmid in nodes_by_pmid
    ]
    elapsed = time.perf_counter() - t0
    _api_log.info("op=efetch parsed=%d duration=%.2fs", len(articles), elapsed)
    return articles


# ── PMC full-text helpers ─────────────────────────────────────────────────────

def _extract_pmcid(linkset: dict) -> tuple[str, str] | None:
    ids = linkset.get("ids", [])
    if not ids:
        return None
    pmid = str(ids[0])
    for lsdb in linkset.get("linksetdbs", []):
        if lsdb.get("linkname") == "pubmed_pmc" and lsdb.get("links"):
            return pmid, str(lsdb["links"][0])
    return None

def _fetch_pmcid_batch(batch: list[str]) -> dict[str, str]:
    resp = _ncbi_post(
        f"{cfg.PUBMED_BASE}/elink.fcgi",
        params={"dbfrom": "pubmed", "db": "pmc", "linkname": "pubmed_pmc", "id": batch, "retmode": "json"},
        timeout=cfg.TIMEOUT_ELINK,
    )
    resp.raise_for_status()
    return {
        p[0]: p[1]
        for p in map(_extract_pmcid, resp.json().get("linksets", []))
        if p is not None
    }


def get_pmcids(pmids: list[str]) -> dict[str, str]:
    """Return {pmid: pmcid} for articles that have PMC full text.

    Sends PMIDs in batches via POST to avoid URL-length timeouts on large sets.
    """
    if not pmids:
        return {}
    _api_log.debug("op=elink pmids=%d", len(pmids))
    t0 = time.perf_counter()
    result: dict[str, str] = {}
    for i in range(0, len(pmids), cfg.ELINK_BATCH_SIZE):
        result.update(_fetch_pmcid_batch(pmids[i:i + cfg.ELINK_BATCH_SIZE]))
    elapsed = time.perf_counter() - t0
    _api_log.info("op=elink pmc_hits=%d/%d duration=%.2fs", len(result), len(pmids), elapsed)
    return result


def fetch_pmc_full_text(pmcid: str) -> str:
    """Fetch PMC article XML and return extracted plain text from the body."""
    _api_log.debug("op=pmc_fulltext pmcid=%s", pmcid)
    t0 = time.perf_counter()
    resp = _ncbi_get(
        f"{cfg.PUBMED_BASE}/efetch.fcgi",
        params={"db": "pmc", "id": pmcid, "rettype": "full", "retmode": "xml"},
        timeout=cfg.TIMEOUT_PMC_FULLTEXT,
    )
    resp.raise_for_status()
    root = ET.fromstring(resp.content)

    body = root.find(".//body")
    target = body if body is not None else root

    parts = []
    for elem in target.iter():
        if elem.tag in {"p", "title"}:
            text = " ".join("".join(elem.itertext()).split())
            if len(text) > 20:
                parts.append(text)

    full_text = "\n\n".join(parts)
    elapsed = time.perf_counter() - t0
    _api_log.info(
        "op=pmc_fulltext pmcid=%s status=%d chars=%d paragraphs=%d duration=%.2fs",
        pmcid, resp.status_code, len(full_text), len(parts), elapsed,
    )
    return full_text


# ── Routes ────────────────────────────────────────────────────────────────────

def _clamp_max_results(raw: str) -> int:
    value = int(raw) if raw else cfg.MAX_RESULTS_DEFAULT
    return max(cfg.MAX_RESULTS_MIN, min(cfg.MAX_RESULTS_MAX, (value // cfg.MAX_RESULTS_MIN) * cfg.MAX_RESULTS_MIN))


def _attach_similarities(user_query: str, results: list[dict]):
    query_emb = embed_texts([user_query])[0]
    article_texts = [f"{a['title']}. {a['abstract']}" for a in results]
    article_embs = embed_texts(article_texts)
    for art, emb in zip(results, article_embs):
        art["similarity"] = round(cosine_similarity(query_emb, emb), 3)


def _run_search(user_query: str, max_results: int) -> tuple[list[dict] | None, str | None, str | None]:
    """Return (results, pubmed_query, error)."""
    try:
        pubmed_query = build_pubmed_query(user_query)
        pmids = search_pubmed(pubmed_query, max_results)
        results = fetch_articles(pmids)
        if results:
            _attach_similarities(user_query, results)
        return results, pubmed_query, None
    except litellm.exceptions.AuthenticationError:
        return None, None, "Invalid Anthropic API key. Set ANTHROPIC_API_KEY in your .env file."
    except requests.RequestException as exc:
        return None, None, f"PubMed request failed: {exc}"
    except Exception as exc:
        return None, None, str(exc)


@app.context_processor
def inject_current_user():
    """Makes current_user_email available to every template (None when logged out)."""
    if getattr(g, "user_id", None) is None:
        return {"current_user_email": None}
    user = db.get_user_by_id(g.user_id)
    return {"current_user_email": user["email"] if user else None}


def _issue_auth_cookies(user_id: int) -> Response:
    """Set the access + refresh cookies for a freshly authenticated user_id."""
    access_token = auth.create_access_token(user_id)
    refresh_token = secrets.token_urlsafe(32)
    refresh_expires = datetime.datetime.now(datetime.timezone.utc) + datetime.timedelta(
        days=cfg.JWT_REFRESH_TTL_DAYS
    )
    db.create_refresh_token(user_id, auth.hash_token(refresh_token), refresh_expires)

    resp = make_response(jsonify({"ok": True}))
    resp.set_cookie(
        auth.ACCESS_COOKIE_NAME, access_token,
        httponly=True, secure=cfg.COOKIE_SECURE, samesite="Lax",
        max_age=cfg.JWT_ACCESS_TTL_MINUTES * 60,
    )
    # Scoped to /auth so it's only ever sent to the refresh/logout endpoints, not on
    # every request the way the access cookie is.
    resp.set_cookie(
        auth.REFRESH_COOKIE_NAME, refresh_token,
        httponly=True, secure=cfg.COOKIE_SECURE, samesite="Lax",
        max_age=cfg.JWT_REFRESH_TTL_DAYS * 86400, path="/auth",
    )
    return resp


def _send_verification(user_id: int, email: str) -> None:
    """Create a fresh email-verification token and email the confirmation link."""
    raw_token = secrets.token_urlsafe(32)
    expires = datetime.datetime.now(datetime.timezone.utc) + datetime.timedelta(
        minutes=cfg.EMAIL_VERIFICATION_TTL_MINUTES
    )
    db.create_email_verification_token(user_id, auth.hash_token(raw_token), expires)
    auth.send_verification_email(email, raw_token)


@app.route("/login", methods=["GET"])
def login_page():
    return render_template("login.html")


@app.route("/register", methods=["GET"])
def register_page():
    return render_template("register.html")


@app.route("/verify-email", methods=["GET"])
def verify_email_page():
    raw_token = request.args.get("token", "")
    valid = db.get_valid_email_verification_token(auth.hash_token(raw_token)) if raw_token else None
    if not valid:
        return render_template("verify_email.html", success=False)

    db.mark_user_verified(valid["user_id"])
    db.mark_email_verification_token_used(valid["id"])
    return render_template("verify_email.html", success=True)


@app.route("/auth/register", methods=["POST"])
@limiter.limit(cfg.REGISTER_RATE_LIMIT)
def auth_register():
    data = request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()
    password = data.get("password") or ""

    if not email or "@" not in email:
        return jsonify({"error": "A valid email is required."}), 400
    if len(password) < 8:
        return jsonify({"error": "Password must be at least 8 characters."}), 400
    if db.get_user_by_email(email):
        return jsonify({"error": "An account with that email already exists."}), 409

    user_id = db.create_user(email, auth.hash_password(password))
    _send_verification(user_id, email)
    return jsonify({
        "ok": True,
        "message": "Check your email for a link to verify your account before logging in.",
    })


@app.route("/auth/login", methods=["POST"])
@limiter.limit(cfg.LOGIN_RATE_LIMIT)
def auth_login():
    data = request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()
    password = data.get("password") or ""

    user = db.get_user_by_email(email)
    if not user or not user["is_active"] or not auth.verify_password(password, user["password_hash"]):
        return jsonify({"error": "Invalid email or password."}), 401
    if not user["email_verified"]:
        return jsonify({
            "error": "Please verify your email before logging in.",
            "code": "email_not_verified",
        }), 403

    return _issue_auth_cookies(user["id"])


@app.route("/auth/resend-verification", methods=["POST"])
@limiter.limit(cfg.RESEND_VERIFICATION_RATE_LIMIT)
def auth_resend_verification():
    data = request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()

    # Always return the same generic response whether or not the account exists/is
    # already verified — otherwise this endpoint becomes an email-enumeration oracle.
    generic_response = jsonify({
        "ok": True,
        "message": "If an unverified account with that email exists, a new link has been sent.",
    })

    user = db.get_user_by_email(email) if email else None
    if user and user["is_active"] and not user["email_verified"]:
        _send_verification(user["id"], user["email"])

    return generic_response


@app.route("/auth/refresh", methods=["POST"])
def auth_refresh():
    raw_token = request.cookies.get(auth.REFRESH_COOKIE_NAME)
    if not raw_token:
        return jsonify({"error": "Not authenticated."}), 401

    token_hash = auth.hash_token(raw_token)
    valid = db.get_valid_refresh_token(token_hash)
    if not valid:
        return jsonify({"error": "Not authenticated."}), 401

    # Rotate: the old refresh token is single-use — revoke it, issue a fresh pair.
    db.revoke_refresh_token(token_hash)
    return _issue_auth_cookies(valid["user_id"])


@app.route("/auth/logout", methods=["POST"])
def auth_logout():
    raw_token = request.cookies.get(auth.REFRESH_COOKIE_NAME)
    if raw_token:
        db.revoke_refresh_token(auth.hash_token(raw_token))
    resp = make_response(jsonify({"ok": True}))
    resp.delete_cookie(auth.ACCESS_COOKIE_NAME)
    resp.delete_cookie(auth.REFRESH_COOKIE_NAME, path="/auth")
    return resp


@app.route("/forgot-password", methods=["GET"])
def forgot_password_page():
    return render_template("forgot_password.html")


@app.route("/reset-password", methods=["GET"])
def reset_password_page():
    return render_template("reset_password.html")


@app.route("/auth/forgot-password", methods=["POST"])
@limiter.limit(cfg.FORGOT_PASSWORD_RATE_LIMIT)
def auth_forgot_password():
    data = request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()

    # Always return the same generic response whether or not the account exists —
    # otherwise this endpoint becomes an email-enumeration oracle.
    generic_response = jsonify({
        "ok": True,
        "message": "If an account with that email exists, a reset link has been sent.",
    })

    user = db.get_user_by_email(email) if email else None
    if user and user["is_active"]:
        raw_token = secrets.token_urlsafe(32)
        expires = datetime.datetime.now(datetime.timezone.utc) + datetime.timedelta(
            minutes=cfg.PASSWORD_RESET_TTL_MINUTES
        )
        db.create_password_reset_token(user["id"], auth.hash_token(raw_token), expires)
        auth.send_password_reset_email(user["email"], raw_token)

    return generic_response


@app.route("/auth/reset-password", methods=["POST"])
def auth_reset_password():
    data = request.get_json(silent=True) or {}
    raw_token = data.get("token") or ""
    new_password = data.get("password") or ""

    if len(new_password) < 8:
        return jsonify({"error": "Password must be at least 8 characters."}), 400

    valid = db.get_valid_password_reset_token(auth.hash_token(raw_token))
    if not valid:
        return jsonify({"error": "This reset link is invalid or has expired."}), 400

    db.update_user_password(valid["user_id"], auth.hash_password(new_password))
    db.mark_password_reset_token_used(valid["id"])
    # A password reset is a "log out everywhere" signal — any stolen/leaked refresh
    # token from before the reset should stop working.
    db.revoke_all_refresh_tokens(valid["user_id"])

    return jsonify({"ok": True})


@app.route("/", methods=["GET", "POST"])
@auth.login_required_page
def index():
    results = None
    pubmed_query = None
    error = None
    user_query = ""
    max_results = cfg.MAX_RESULTS_DEFAULT

    if request.method == "POST":
        user_query = request.form.get("query", "").strip()
        if user_query:
            max_results = _clamp_max_results(request.form.get("max_results", str(cfg.MAX_RESULTS_DEFAULT)))
            results, pubmed_query, error = _run_search(user_query, max_results)

    return render_template(
        "index.html",
        user_query=user_query,
        pubmed_query=pubmed_query,
        results=results,
        error=error,
        max_results=max_results,
    )


@app.route("/collections", methods=["GET"])
@auth.login_required_page
def collections():
    items = db.list_collections(g.user_id)
    return render_template("collections.html", collections=items)


@app.route("/collections", methods=["POST"])
@auth.login_required
def collections_save():
    """Receive selected articles, fetch PMC full text where available, chunk, embed, store."""
    payload = request.get_json(force=True)
    name = (payload.get("name") or "").strip()
    user_query = payload.get("user_query", "")
    pubmed_query = payload.get("pubmed_query", "")
    articles = payload.get("articles", [])

    if not name:
        return {"error": "Collection name is required."}, 400
    if not articles:
        return {"error": "No articles to save."}, 400

    try:
        pmids = [a["pmid"] for a in articles]
        pmcid_map = get_pmcids(pmids)
        cid = db.create_collection(name, user_query, pubmed_query, g.user_id)

        # Phase 1: fetch PMC full texts in parallel (I/O-bound network calls)
        max_workers = min(len(articles), 8)
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            fetched = list(executor.map(_fetch_article_text, articles, [pmcid_map] * len(articles)))

        # Phase 2: persist article rows; collect chunks that still need embedding
        pending: list[tuple[str, list[str]]] = []  # [(pmid, chunks), ...]
        for art, full_text, pmcid in fetched:
            db.add_article(cid, art, has_full_text=bool(full_text), pmcid=pmcid)
            pmid = art["pmid"]
            if not db.chunks_exist(pmid):
                text_to_chunk = full_text or f"{art['title']}. {art['abstract']}"
                pending.append((pmid, chunk_text(text_to_chunk, cfg.CHUNK_SIZE, cfg.CHUNK_OVERLAP)))

        # Phase 3: embed all pending chunks in one batch, then save
        if pending:
            all_chunks = [c for _, chunks in pending for c in chunks]
            all_embeddings = embed_texts(all_chunks)
            offset = 0
            for pmid, chunks in pending:
                n = len(chunks)
                db.save_chunks(pmid, chunks, all_embeddings[offset:offset + n])
                offset += n

        return {"id": cid}
    except Exception as exc:
        return {"error": str(exc)}, 500


def _sse(payload: dict) -> str:
    return f"data: {json.dumps(payload)}\n\n"


def _save_fetch_phase(articles: list[dict], pmcid_map: dict):
    """Generator — yields SSE fetch events; returns (fetched, counts) via StopIteration.value."""
    total = len(articles)
    fetched: list[tuple[dict, str | None, str | None]] = []
    counts = {"full_text": 0, "abstract": 0, "fallback": 0}

    with ThreadPoolExecutor(max_workers=min(total, 8)) as executor:
        futures = {executor.submit(_fetch_article_text, art, pmcid_map): art for art in articles}
        for future in as_completed(futures):
            art, full_text, pmcid = future.result()
            fetched.append((art, full_text, pmcid))
            had_pmcid = bool(pmcid_map.get(futures[future]["pmid"]))
            if full_text:
                counts["full_text"] += 1
            elif had_pmcid:
                counts["fallback"] += 1
            else:
                counts["abstract"] += 1
            yield _sse({"type": "fetch", "done": len(fetched), "total": total, **counts})

    return fetched, counts


def _save_persist_phase(cid: int, fetched: list) -> list[tuple[str, list[str], bool]]:
    """Insert article rows; return (pmid, chunks, is_full_text) triples that still need embedding."""
    pending: list[tuple[str, list[str], bool]] = []
    for art, full_text, pmcid in fetched:
        db.add_article(cid, art, has_full_text=bool(full_text), pmcid=pmcid)
        pmid = art["pmid"]
        if not db.chunks_exist(pmid):
            text_to_chunk = full_text or f"{art['title']}. {art['abstract']}"
            chunks = chunk_text(text_to_chunk, cfg.CHUNK_SIZE, cfg.CHUNK_OVERLAP)
            pending.append((pmid, chunks, bool(full_text)))
    return pending


def _save_embed_phase(pending: list[tuple[str, list[str], bool]]):
    """Generator — yields SSE embedding event then embeds and saves all pending chunks."""
    if not pending:
        return
    full_text_count = sum(1 for _, _, is_full_text in pending if is_full_text)
    yield _sse({"type": "embedding", "count": len(pending),
                "full_text": full_text_count,
                "abstract": len(pending) - full_text_count})
    all_chunks = [c for _, chunks, _ in pending for c in chunks]
    all_embeddings = embed_texts(all_chunks)
    offset = 0
    for pmid, chunks, _ in pending:
        n = len(chunks)
        db.save_chunks(pmid, chunks, all_embeddings[offset:offset + n])
        offset += n


@app.route("/collections/save-stream", methods=["POST"])
@auth.login_required
def collections_save_stream():
    """SSE endpoint: save collection with per-article fetch + embedding progress."""
    payload = request.get_json(force=True)
    name = (payload.get("name") or "").strip()
    user_query = payload.get("user_query", "")
    pubmed_query = payload.get("pubmed_query", "")
    articles = payload.get("articles", [])
    user_id = g.user_id

    if not name:
        return {"error": "Collection name is required."}, 400
    if not articles:
        return {"error": "No articles to save."}, 400

    def generate():
        try:
            pmcid_map = get_pmcids([a["pmid"] for a in articles])
            cid = db.create_collection(name, user_query, pubmed_query, user_id)
            fetched, counts = yield from _save_fetch_phase(articles, pmcid_map)
            pending = _save_persist_phase(cid, fetched)
            yield from _save_embed_phase(pending)
            yield _sse({"type": "done", "id": cid, "count": len(articles), **counts})
        except Exception as exc:
            yield _sse({"type": "error", "message": str(exc)})

    return Response(
        stream_with_context(generate()),
        mimetype=_SSE_MIMETYPE,
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


def _try_fetch_full_text(pmcid: str) -> tuple[str | None, str | None]:
    """Return (full_text, pmcid) — pmcid is set to None on fetch failure."""
    try:
        return fetch_pmc_full_text(pmcid), pmcid
    except Exception:
        return None, None


def _fetch_article_text(art: dict, pmcid_map: dict) -> tuple[dict, str | None, str | None]:
    """Resolve PMC full text for one article (called in parallel).

    Returns (art, full_text, resolved_pmcid).
    """
    pmcid = pmcid_map.get(art["pmid"])
    if pmcid:
        full_text, resolved_pmcid = _try_fetch_full_text(pmcid)
    else:
        full_text, resolved_pmcid = None, None
    return art, full_text, resolved_pmcid


def generate_starter_questions(user_query: str, articles: list[dict]) -> list[str]:
    """Generate 4 opening questions for a collection using its topic and article titles."""
    titles = "\n".join(f"- {a['title']}" for a in articles[:10])
    _claude_log.debug("op=starter_questions topic=%r articles=%d", user_query[:80], len(articles))
    t0 = time.perf_counter()
    try:
        kwargs = {}
        if cfg.STARTER_QUESTIONS_MODEL.startswith("ollama/"):
            kwargs["api_base"] = cfg.OLLAMA_BASE_URL
        response = litellm.completion(
            model=cfg.STARTER_QUESTIONS_MODEL,
            max_tokens=cfg.MAX_TOKENS_STARTER_QS,
            messages=[
                {"role": "system", "content": cfg.PROMPT_STARTER_QUESTIONS},
                {"role": "user", "content": f"Topic: {user_query}\n\nArticles:\n{titles}"},
            ],
            **kwargs,
        )
        content = (response.choices[0].message.content or "").strip()
        questions = json.loads(_strip_json_fence(content))[:4]
        elapsed = time.perf_counter() - t0
        usage = response.usage
        _claude_log.info(
            "op=starter_questions in_tokens=%d out_tokens=%d questions=%d duration=%.2fs",
            getattr(usage, "prompt_tokens", 0), getattr(usage, "completion_tokens", 0), len(questions), elapsed,
        )
        return questions
    except Exception as exc:
        _claude_log.warning("op=starter_questions failed: %s", exc)
        return []


@app.route("/collections/<int:cid>", methods=["GET"])
@auth.login_required_page
def collection_detail(cid: int):
    collection = db.get_collection(cid, g.user_id)
    if not collection:
        return "Collection not found", 404
    articles = db.get_collection_articles(cid)
    return render_template(
        "collection.html",
        collection=collection,
        articles=articles,
    )


@app.route("/collections/<int:cid>/export.csv", methods=["GET"])
@auth.login_required
def collection_export_csv(cid: int):
    collection = db.get_collection(cid, g.user_id)
    if not collection:
        return "Collection not found", 404
    articles = db.get_collection_articles(cid)

    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["PMID", "Journal", "Title", "Year", "Authors", "Abstract"])
    for art in articles:
        writer.writerow([
            art["pmid"],
            art["journal"] or "",
            art["title"] or "",
            art["year"] or "",
            art["authors"] or "",
            art["abstract"] or "",
        ])

    safe_name = re.sub(r'[^\w\s-]', '', collection["name"]).strip().replace(' ', '-') or "collection"
    return Response(
        buf.getvalue().encode("utf-8-sig"),
        mimetype="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.csv"'},
    )


@app.route("/collections/<int:cid>/starter-questions", methods=["GET"])
@auth.login_required
def collection_starter_questions(cid: int):
    # Ownership must be checked before the cache lookup, not inside the cache-miss
    # branch — otherwise a cache hit would skip authorization and leak another
    # user's cached questions to anyone who guesses their cid.
    collection = db.get_collection(cid, g.user_id)
    if not collection:
        return {"questions": []}
    if cid not in _starter_questions_cache:
        articles = db.get_collection_articles(cid)
        _starter_questions_cache[cid] = generate_starter_questions(
            collection["user_query"], articles
        )
    return {"questions": _starter_questions_cache[cid]}



def _sse_emit(payload: dict) -> str:
    return f"data: {json.dumps(payload)}\n\n"


def _strip_json_fence(text: str) -> str:
    """Strip a ```json ... ``` (or bare ```) markdown fence some models wrap JSON output in."""
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text)
    return text.strip()


def _parse_suggestions(raw: str) -> list[str]:
    try:
        return json.loads(_strip_json_fence(raw))[:4]
    except Exception:
        return []


def _compute_cost(model: str, input_tokens: int, output_tokens: int) -> float:
    """Return the USD cost for one API call at provider list prices."""
    pricing = cfg.CHAT_MODELS.get(model, {"input_price": 0.0, "output_price": 0.0})
    return (input_tokens * pricing["input_price"] + output_tokens * pricing["output_price"]) / 1_000_000


def _stream_delimited_response(
    model: str, context: str, question: str, delim: str, emit
) -> tuple[list[str], str, dict]:
    """Stream a Claude response split by *delim*.

    Returns (sse_events, post_delimiter_text, usage_dict).
    usage_dict has keys: input_tokens, output_tokens.
    """
    pre = ""
    post = ""
    found = False
    events: list[str] = []

    kwargs = {}
    if model.startswith("ollama_chat/"):
        kwargs["api_base"] = cfg.OLLAMA_CLOUD_BASE_URL
        kwargs["api_key"] = cfg.OLLAMA_API_KEY
    elif model.startswith("ollama/"):
        kwargs["api_base"] = cfg.OLLAMA_BASE_URL

    usage = {"input_tokens": 0, "output_tokens": 0}
    stream = litellm.completion(
        model=model,
        max_tokens=cfg.MAX_TOKENS_RAG_RESPONSE,
        stream=True,
        stream_options={"include_usage": True},
        messages=[
            {"role": "system", "content": cfg.PROMPT_RAG_SYSTEM},
            {"role": "user", "content": f"Articles:\n\n{context}\n\nQuestion: {question}"},
        ],
        **kwargs,
    )
    for chunk in stream:
        chunk_usage = getattr(chunk, "usage", None)
        if chunk_usage is not None:
            usage = {
                "input_tokens":  getattr(chunk_usage, "prompt_tokens", 0) or 0,
                "output_tokens": getattr(chunk_usage, "completion_tokens", 0) or 0,
            }
        if not chunk.choices:
            continue
        text = chunk.choices[0].delta.content or ""
        if not text:
            continue
        pre, post, found, chunk_events = _advance_delimiter_state(
            pre, post, found, text, delim, emit
        )
        events.extend(chunk_events)

    if not found and pre:
        events.append(emit({"text": pre}))

    return events, post, usage


def _advance_delimiter_state(
    pre: str, post: str, found: bool, text: str, delim: str, emit
) -> tuple[str, str, bool, list[str]]:
    """Process one streamed token; return updated (pre, post, found, new_events)."""
    events: list[str] = []
    if found:
        return pre, post + text, found, events

    pre += text
    if delim in pre:
        idx = pre.index(delim)
        safe = pre[:idx]
        post = pre[idx + len(delim):]
        pre = ""
        if safe:
            events.append(emit({"text": safe}))
        return pre, post, True, events

    safe_len = max(0, len(pre) - len(delim) + 1)
    if safe_len:
        events.append(emit({"text": pre[:safe_len]}))
        pre = pre[safe_len:]
    return pre, post, found, events


def _build_rag_context(top_chunks: list[dict]) -> tuple[str, list[dict]]:
    """Return (context_string, deduplicated_citations).

    Each unique article is assigned one citation number. Chunks from the same
    article share that number so Claude's inline [N] markers always match the
    numbers shown in the Sources section.
    """
    context_parts = []
    pmid_to_num: dict[str, int] = {}
    citations = []
    for chunk in top_chunks:
        pmid = chunk["pmid"]
        if pmid not in pmid_to_num:
            pmid_to_num[pmid] = len(citations) + 1
            citations.append(chunk)
        num = pmid_to_num[pmid]
        context_parts.append(
            f"[{num}] {chunk['title']}\n"
            f"Authors: {chunk['authors']}\n"
            f"Journal: {chunk['journal']} ({chunk['year']})\n"
            f"Excerpt: {chunk['chunk_text']}"
        )
    return "\n\n".join(context_parts), citations


def _extract_answer_from_events(events: list[str]) -> str:
    """Reconstruct plain answer text from SSE event strings."""
    parts = []
    for event in events:
        if not event.startswith("data: "):
            continue
        try:
            payload = json.loads(event[6:].rstrip("\n"))
            if "text" in payload:
                parts.append(payload["text"])
        except Exception:
            pass
    return "".join(parts)


_CONVERSATION_NOT_FOUND = "Conversation not found."


def _get_owned_conversation(vid: int) -> dict | None:
    """Return the conversation only if it exists AND its collection is owned by g.user_id."""
    conv = db.get_conversation(vid)
    if not conv or conv["user_id"] != g.user_id:
        return None
    return conv


@app.route("/collections/<int:cid>/conversations", methods=["GET"])
@auth.login_required
def collection_conversations(cid: int):
    if not db.get_collection(cid, g.user_id):
        return {"error": "Collection not found."}, 404
    return db.list_conversations(cid)


@app.route("/conversations/<int:vid>/messages", methods=["GET"])
@auth.login_required
def conversation_messages(vid: int):
    if not _get_owned_conversation(vid):
        return {"error": _CONVERSATION_NOT_FOUND}, 404
    return db.get_conversation_messages(vid)


@app.route("/conversations/<int:vid>/rename", methods=["PATCH"])
@auth.login_required
def conversation_rename(vid: int):
    if not _get_owned_conversation(vid):
        return {"error": _CONVERSATION_NOT_FOUND}, 404
    data  = request.get_json(force=True)
    title = (data.get("title") or "").strip()
    if not title:
        return {"error": "Title cannot be empty."}, 400
    db.rename_conversation(vid, title)
    return {"ok": True}


@app.route("/conversations/<int:vid>/delete", methods=["POST"])
@auth.login_required
def conversation_delete(vid: int):
    if not _get_owned_conversation(vid):
        return {"error": _CONVERSATION_NOT_FOUND}, 404
    db.delete_conversation(vid)
    return {"ok": True}


@app.route("/api/models", methods=["GET"])
@auth.login_required
def api_models():
    """List chat models currently available: env-key-gated providers + discovered Ollama models."""
    models = [
        {"id": mid, "display": meta["display"], "provider": meta["provider"]}
        for mid, meta in available_chat_models().items()
    ]
    return {"models": models, "default": cfg.DEFAULT_CHAT_MODEL}


@app.route("/collections/<int:cid>/ask", methods=["POST"])
@auth.login_required
def collection_ask(cid: int):
    """SSE endpoint: embeds question, retrieves top-k chunks, streams RAG answer."""
    if not db.get_collection(cid, g.user_id):
        return {"error": "Collection not found."}, 404
    data = request.get_json(force=True)
    question = (data.get("question") or "").strip()
    if not question:
        return {"error": "No question provided."}, 400
    model = data.get("model", cfg.DEFAULT_CHAT_MODEL)
    if model not in available_chat_models():
        model = cfg.DEFAULT_CHAT_MODEL
    conversation_id = data.get("conversation_id")
    # A client-supplied conversation_id could belong to another user's conversation
    # (or a different collection); only trust it if it's actually owned by this user
    # and attached to this collection, otherwise start a fresh one.
    if conversation_id is not None:
        owned_conv = _get_owned_conversation(conversation_id)
        if not owned_conv or owned_conv["collection_id"] != cid:
            conversation_id = None

    q_emb = embed_texts([question])[0]
    top_chunks = db.semantic_search(cid, q_emb, k=cfg.SEMANTIC_SEARCH_K)

    if not top_chunks:
        # Still persist the unanswered question so the conversation exists
        if conversation_id is None:
            conversation_id = db.create_conversation(cid, question)
        db.add_message(conversation_id, "user", question)
        db.add_message(conversation_id, "assistant", "No articles found in this collection.")

        def empty():
            yield f"data: {json.dumps({'text': 'No articles found in this collection.'})}\n\n"
            yield f"data: {json.dumps({'done': True, 'citations': [], 'conversation_id': conversation_id})}\n\n"
        return Response(empty(), mimetype=_SSE_MIMETYPE,
                        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

    # Create / continue conversation
    if conversation_id is None:
        conversation_id = db.create_conversation(cid, question)
    db.add_message(conversation_id, "user", question)

    context, citations = _build_rag_context(top_chunks)

    # Slim, serialisable form — num field lets the frontend map [N] → URL
    stored_citations = [
        {
            "num": i + 1,
            "pmid": c["pmid"],
            "title": c["title"],
            "url": c["url"],
            "journal": c["journal"],
            "year": c["year"],
            "similarity": c["similarity"],
        }
        for i, c in enumerate(citations)
    ]

    _DELIM = cfg.RAG_DELIMITER

    def generate():
        t0 = time.perf_counter()
        _claude_log.debug(
            "op=collection_ask cid=%d chunks=%d question=%r",
            cid, len(top_chunks), question[:120],
        )
        try:
            answer_text, suggestions_json, usage = _stream_delimited_response(
                model, context, question, _DELIM, _sse_emit
            )
        except Exception as exc:
            _claude_log.exception("op=collection_ask FAILED cid=%d: %s", cid, exc)
            yield _sse_emit({"error": str(exc)})
            return

        yield from answer_text

        full_answer = _extract_answer_from_events(answer_text)
        db.add_message(conversation_id, "assistant", full_answer, stored_citations)

        suggestions = _parse_suggestions(suggestions_json)
        elapsed = time.perf_counter() - t0
        cost = _compute_cost(model, usage["input_tokens"], usage["output_tokens"])
        usage_info = {
            "model":         model,
            "input_tokens":  usage["input_tokens"],
            "output_tokens": usage["output_tokens"],
            "cost_usd":      round(cost, 6),
        }
        _claude_log.info(
            "op=collection_ask cid=%d vid=%d citations=%d suggestions=%d "
            "in_tok=%d out_tok=%d cost_usd=%.6f duration=%.2fs",
            cid, conversation_id, len(stored_citations), len(suggestions),
            usage["input_tokens"], usage["output_tokens"], cost, elapsed,
        )
        yield f"data: {json.dumps({'done': True, 'citations': stored_citations, 'suggestions': suggestions, 'conversation_id': conversation_id, 'usage': usage_info})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype=_SSE_MIMETYPE,
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/collections/<int:cid>/delete", methods=["POST"])
@auth.login_required
def collection_delete(cid: int):
    db.delete_collection(cid, g.user_id)
    return {"ok": True}


# ── Conversation export ───────────────────────────────────────────────────────

def _rtf_esc(text: str) -> str:
    """Escape a string for inclusion in an RTF document body."""
    out = []
    for ch in text:
        if ch == '\\':
            out.append('\\\\')
        elif ch == '{':
            out.append('\\{')
        elif ch == '}':
            out.append('\\}')
        elif ord(ch) > 127:
            out.append(f'\\u{ord(ch)}?')
        else:
            out.append(ch)
    return ''.join(out)


def _rtf_hyperlink(url: str, text: str) -> str:
    """Return an RTF field that renders *text* as a clickable hyperlink to *url*."""
    safe_url = url.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
    return (
        r'{\field{\*\fldinst{HYPERLINK "' + safe_url + r'"}}' +
        r'{\fldrslt \cf2\ul ' + _rtf_esc(text) + r'}}'
    )


def _rtf_line_with_links(text: str, url_map: dict) -> str:
    """Escape *text* for RTF, converting [N] markers to hyperlinks via *url_map*."""
    parts = re.split(r'(\[\d+\])', text)
    out = []
    for part in parts:
        m = re.match(r'^\[(\d+)\]$', part)
        url = m and url_map.get(int(m.group(1)))
        out.append(_rtf_hyperlink(url, part) if url else _rtf_esc(part))
    return ''.join(out)


def _rtf_citation_line(c: dict) -> str:
    """Return the RTF line string for a single citation entry."""
    num     = c.get('num', '')
    url     = c.get('url', '')
    ctitle  = c.get('title', '')
    journal = c.get('journal')
    meta    = ''
    if journal:
        year = c.get('year')
        meta = f" - {journal} ({year})" if year else f" - {journal}"
    link = _rtf_hyperlink(url, ctitle) if url else _rtf_esc(ctitle)
    return r'\pard [' + _rtf_esc(str(num)) + '] ' + link + _rtf_esc(meta) + r'\par'


def _rtf_user_turn(msg: dict) -> str:
    """Return the RTF line string for a user message."""
    return r'\pard\cf1\b You:\b0\cf0 ' + _rtf_esc(msg['content']) + r'\par'


def _rtf_assistant_turn(msg: dict) -> list[str]:
    """Return the RTF line strings for an assistant message with optional citations."""
    citations = msg.get('citations') or []
    url_map   = {c['num']: c['url'] for c in citations if 'num' in c and 'url' in c}
    result    = [r'\pard\b Claude:\b0\par']
    for line in msg['content'].split('\n'):
        result.append(r'\pard ' + _rtf_line_with_links(line, url_map) + r'\par')
    if citations:
        result.append(r'\pard\par\pard\b\fs18 Sources\b0\fs20\par')
        for c in citations:
            result.append(_rtf_citation_line(c))
    return result


def _build_rtf(title: str, collection_name: str, created_at: str, messages: list[dict]) -> str:
    lines = [
        r'{\rtf1\ansi\ansicpg1252\deff0',
        r'{\fonttbl{\f0\fswiss\fcharset0 Arial;}}',
        r'{\colortbl;\red15\green52\blue96;\red5\green99\blue193;}',  # \cf1=navy \cf2=link-blue
        r'\f0\fs20',
        r'\pard\b\fs28 ' + _rtf_esc(title) + r'\b0\fs20\par',
        r'\pard\i Collection: ' + _rtf_esc(collection_name) + r'\i0\par',
        r'\pard\i Date: ' + _rtf_esc(created_at[:10]) + r'\i0\par',
        r'\pard\par',
    ]
    for msg in messages:
        if msg['role'] == 'user':
            lines.append(_rtf_user_turn(msg))
        else:
            lines.extend(_rtf_assistant_turn(msg))
        lines.append(r'\pard\par')
    lines.append('}')
    return '\n'.join(lines)


# ── Docx hyperlink helpers ────────────────────────────────────────────────────

def _docx_add_hyperlink(para, url: str, text: str):
    """Append a clickable hyperlink run to *para*."""
    r_id = para.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    run.append(rpr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hl.append(run)
    para._p.append(hl)


def _docx_append_with_links(para, text: str, url_map: dict):
    """Append *text* to *para*, converting [N] markers to hyperlinks."""
    for part in re.split(r'(\[\d+\])', text):
        m = re.match(r'^\[(\d+)\]$', part)
        url = m and url_map.get(int(m.group(1)))
        if url:
            _docx_add_hyperlink(para, url, part)
        else:
            para.add_run(part)


def _render_citation(doc, c: dict):
    """Append a single citation paragraph to *doc*."""
    item = doc.add_paragraph()
    item.add_run(f'[{c.get("num", "")}] ')
    url    = c.get('url', '')
    ctitle = c.get('title', '')
    if url:
        _docx_add_hyperlink(item, url, ctitle)
    else:
        item.add_run(ctitle)
    journal = c.get('journal')
    if journal:
        year    = c.get('year')
        suffix  = f" \u2013 {journal} ({year})" if year else f" \u2013 {journal}"
        item.add_run(suffix).italic = True


def _render_user_turn(doc, msg: dict):
    """Append a user message paragraph to *doc*."""
    p     = doc.add_paragraph()
    label = p.add_run('You:  ')
    label.bold = True
    label.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
    p.add_run(msg['content'])


def _render_assistant_turn(doc, msg: dict):
    """Append an assistant message (with optional citations) to *doc*."""
    citations = msg.get('citations') or []
    url_map   = {c['num']: c['url'] for c in citations if 'num' in c and 'url' in c}

    lines = msg['content'].split('\n')
    p = doc.add_paragraph()
    p.add_run('Claude:  ').bold = True
    _docx_append_with_links(p, lines[0], url_map)
    for line in lines[1:]:
        p = doc.add_paragraph()
        _docx_append_with_links(p, line, url_map)

    if citations:
        doc.add_paragraph().add_run('Sources:').bold = True
        for c in citations:
            _render_citation(doc, c)


def _build_docx(title: str, collection_name: str, created_at: str, messages: list[dict]) -> io.BytesIO:
    doc = Document()
    doc.add_heading(title, level=1)

    meta = doc.add_paragraph()
    meta.add_run(f'Collection: {collection_name}  |  {created_at[:10]}').italic = True
    doc.add_paragraph()

    for msg in messages:
        if msg['role'] == 'user':
            _render_user_turn(doc, msg)
        else:
            _render_assistant_turn(doc, msg)
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/conversations/<int:vid>/export", methods=["GET"])
@auth.login_required
def conversation_export(vid: int):
    fmt = request.args.get("format", "docx").lower()
    if fmt not in ("docx", "rtf"):
        return {"error": "format must be 'docx' or 'rtf'"}, 400

    conv = _get_owned_conversation(vid)
    if not conv:
        return "Conversation not found", 404

    messages = db.get_conversation_messages(vid)
    if not messages:
        return "No messages to export", 404

    safe_name = (
        "".join(c for c in conv['title'] if c.isalnum() or c in ' -_')[:60].strip()
        or "conversation"
    )

    if fmt == "docx":
        buf = _build_docx(conv['title'], conv['collection_name'], conv['created_at'], messages)
        return Response(
            buf.read(),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
        )

    rtf_text = _build_rtf(conv['title'], conv['collection_name'], conv['created_at'], messages)
    return Response(
        rtf_text.encode("ascii", errors="replace"),
        mimetype="application/rtf",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.rtf"'},
    )


# ── Startup ───────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    db.init_db()
    app.run(host="127.0.0.1", debug=True, port=8080)